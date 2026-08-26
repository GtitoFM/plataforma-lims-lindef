openpyxl = None
PatternFill = None

try:
    import openpyxl
    from openpyxl.styles import PatternFill
except ModuleNotFoundError:
    print("⚠️ Módulo openpyxl não encontrado. Algumas funcionalidades de Excel ficarão limitadas.")

Document = None
Inches = None
Pt = None
RGBColor = None
docx = None

try:
    from docx import Document as _DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    import docx.enum.text
    Document = _DocxDocument
except ModuleNotFoundError:
    print("⚠️ Módulo python-docx não encontrado. A geração de documentos será feita em texto simples.")

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime
from typing import List, Dict, Optional, Tuple, Any
import json
import re
import csv
import os
import zipfile
import xml.etree.ElementTree as ET
import sqlite3
import sys
from contextlib import contextmanager
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from urllib.parse import urlparse

try:
    import pandas as pd
except ModuleNotFoundError:
    pd = None
    print("⚠️ Módulo pandas não encontrado. A importação de dados usará leitura nativa.")


class _DocumentoTextoSimples:
    """Fallback simples para gerar documentos sem python-docx."""
    def __init__(self):
        self._conteudo = []

    def add_heading(self, text: str, level: int = 0):
        prefixo = "#" * max(1, level + 1)
        self._conteudo.append(f"{prefixo} {text}")

    def add_paragraph(self, text: str):
        self._conteudo.append(text)

    def save(self, path: str):
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(self._conteudo))


def DocumentFactory(*args, **kwargs):
    if Document is not None:
        return Document(*args, **kwargs)
    return _DocumentoTextoSimples()


# Ajusta o nome global para manter compatibilidade com o restante do código.
Document = DocumentFactory


def _parse_cell_value(cell, shared_strings: List[str]):
    """Extrai o valor de uma célula XLSX sem depender de openpyxl."""
    cell_type = cell.attrib.get('t')
    if cell_type == 's':
        index_text = cell.findtext('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v', default='0')
        try:
            index = int(index_text)
        except ValueError:
            return ''
        return shared_strings[index] if 0 <= index < len(shared_strings) else ''

    if cell_type == 'inlineStr':
        return ''.join(node.text or '' for node in cell.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t'))

    if cell_type == 'b':
        return cell.findtext('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v', default='0') == '1'

    value = cell.findtext('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v', default='')
    if value == '':
        return ''

    try:
        if '.' in value:
            return float(value)
        return int(value)
    except ValueError:
        return value


def _ler_planilha_xlsx(arquivo_excel: str) -> List[List[Any]]:
    """Lê uma planilha .xlsx usando apenas a biblioteca padrão do Python."""
    if not os.path.exists(arquivo_excel):
        raise FileNotFoundError(f'Arquivo não encontrado: {arquivo_excel}')

    with zipfile.ZipFile(arquivo_excel) as zf:
        shared_strings = []
        if 'xl/sharedStrings.xml' in zf.namelist():
            tree = ET.fromstring(zf.read('xl/sharedStrings.xml'))
            for si in tree.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}si'):
                texto = ''.join(node.text or '' for node in si.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t'))
                shared_strings.append(texto)

        workbook = ET.fromstring(zf.read('xl/workbook.xml'))
        rels = ET.fromstring(zf.read('xl/_rels/workbook.xml.rels'))
        rel_map = {rel.attrib['Id']: rel.attrib['Target'] for rel in rels.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')}

        sheets = workbook.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}sheets')
        if sheets is None or len(sheets) == 0:
            return []

        first_sheet = sheets[0]
        rel_id = first_sheet.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']
        target = rel_map.get(rel_id, '')
        sheet_path = target if target.startswith('xl/') else f'xl/{target}'
        if sheet_path.startswith('/'):
            sheet_path = sheet_path.lstrip('/')

        sheet_xml = ET.fromstring(zf.read(sheet_path))
        sheet_data = sheet_xml.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}sheetData')
        if sheet_data is None:
            return []

        linhas = []
        for row in sheet_data.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row'):
            linha = []
            for cell in row.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                linha.append(_parse_cell_value(cell, shared_strings))
            linhas.append(linha)
        return linhas


class Equipamento:
    """Classe base para equipamentos do laboratório com rastreabilidade completa"""
    def __init__(self, nome: str, fabricante: str, status_operacional: str = "Ativo",
                 ultima_calibracao: str = "", pesquisador_responsavel: str = ""):
        self.nome = nome
        self.fabricante = fabricante
        self.status_operacional = status_operacional  
        self.status_disponibilidade = status_operacional
        self.ultima_calibracao = ultima_calibracao or datetime.now().strftime('%d/%m/%Y')
        self.pesquisador_responsavel = pesquisador_responsavel
        self.historico_uso = []  
        self.dados_medicao = []  
        self.documentacao = []  
        self.metas_vinculadas = []  
        self.log_manutencao = []
        self.treino_confiabilidade_horas = 0.0
        self.documentos_tecnicos = []

    def registrar_uso(self, pesquisador: str, proposito: str, data: str = ""):
        """Registra uso do equipamento"""
        uso = {
            'data': data or datetime.now().strftime('%d/%m/%Y %H:%M'),
            'pesquisador': pesquisador,
            'proposito': proposito
        }
        self.historico_uso.append(uso)

    def adicionar_dado_medicao(self, dados: Dict):
        """Adiciona dados de medição capturados automaticamente"""
        dados['timestamp'] = datetime.now().isoformat()
        self.dados_medicao.append(dados)

    def processar_dados_brutos(self, dados_brutos: Any) -> List[Dict]:
        """Converte dados crus em uma estrutura uniforme para o LIMS"""
        if isinstance(dados_brutos, dict):
            registros = [dados_brutos]
        elif isinstance(dados_brutos, list):
            registros = dados_brutos
        elif pd is not None and hasattr(dados_brutos, 'to_dict'):
            registros = dados_brutos.to_dict(orient='records')
        else:
            registros = [{"valor_bruto": dados_brutos}]

        processados = []
        for item in registros:
            if isinstance(item, dict):
                registro = dict(item)
                registro.setdefault('equipamento', self.nome)
                registro.setdefault('processado_em', datetime.now().isoformat())
                processados.append(registro)
            else:
                processados.append({'valor_bruto': item, 'equipamento': self.nome})
        self.dados_medicao.extend(processados)
        return processados

    def validar_manutencao(self, descricao: str, status: str = "Em Manutenção", data: str = "") -> Dict:
        """Registra e valida a disponibilidade de manutenção de um equipamento"""
        if not descricao or not descricao.strip():
            raise ValueError('Descrição da manutenção é obrigatória')
        registro = {
            'descricao': descricao,
            'status': status,
            'data': data or datetime.now().strftime('%d/%m/%Y'),
            'equipamento': self.nome
        }
        self.log_manutencao.append(registro)
        self.status_operacional = status
        self.status_disponibilidade = status
        return registro

    def registrar_treino_confiabilidade(self, horas_pratica: float, observacao: str = "") -> Dict:
        """Registra horas de treino de confiabilidade antes da coleta real"""
        self.treino_confiabilidade_horas = horas_pratica
        return {
            'equipamento': self.nome,
            'horas_pratica': horas_pratica,
            'observacao': observacao,
            'data': datetime.now().strftime('%d/%m/%Y')
        }

    def vincular_meta(self, meta_id: str, pesquisador: str):
        """Vincula uma meta ao equipamento"""
        vinculo = {
            'meta_id': meta_id,
            'pesquisador': pesquisador,
            'data_vinculo': datetime.now().strftime('%d/%m/%Y')
        }
        self.metas_vinculadas.append(vinculo)

    def to_dict(self) -> Dict:
        """Converte equipamento para dicionário"""
        return {
            'nome': self.nome,
            'fabricante': self.fabricante,
            'status_operacional': self.status_operacional,
            'status_disponibilidade': self.status_disponibilidade,
            'ultima_calibracao': self.ultima_calibracao,
            'pesquisador_responsavel': self.pesquisador_responsavel,
            'historico_uso': self.historico_uso,
            'dados_medicao': self.dados_medicao,
            'documentacao': self.documentacao,
            'metas_vinculadas': self.metas_vinculadas,
            'log_manutencao': self.log_manutencao,
            'treino_confiabilidade_horas': self.treino_confiabilidade_horas,
            'documentos_tecnicos': self.documentos_tecnicos
        }


class RDA(Equipamento):
    """Respirador Digital Aberto - Vinculado a Emanuel Fernandes Ferreira da Silva Júnior"""
    def __init__(self, fabricante: str = "Fabricante Padrão", pesquisador_responsavel: str = "Emanuel Fernandes Ferreira da Silva Júnior"):
        super().__init__("RDA", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.problemas_engenharia = []  # Log de problemas de engenharia
        self.log_falhas = []
        self.ficha_avaliativa_status = "90% concluído"

    def registrar_calibracao_seringa(self, volume: float, erro_medido: float, data: str = "", meta_alcance: float = 80.0):
        """Registra calibração com seringa e marca meta de alcance"""
        calibracao = {
            'tipo': 'Calibração com Seringa',
            'volume': volume,
            'erro_medido': erro_medido,
            'meta_alcance': meta_alcance,
            'data': data or datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.adicionar_dado_medicao(calibracao)
        return calibracao

    def calibrar_seringa(self, volume: float, erro_medido: float, data: str = "") -> Dict:
        """Método específico para calibrar a seringa do RDA"""
        return self.registrar_calibracao_seringa(volume, erro_medido, data=data, meta_alcance=80.0)

    def gerar_ficha_avaliativa(self, status: str = "90% concluído", permitir_edicoes: bool = True) -> Dict:
        """Gera ficha avaliativa com status e flag de edição"""
        self.ficha_avaliativa_status = status
        return {
            'equipamento': self.nome,
            'status': status,
            'permitir_edicoes': permitir_edicoes,
            'data': datetime.now().strftime('%d/%m/%Y')
        }

    def registrar_falha(self, descricao: str, severidade: str = "Baixa"):
        """Registra problemas de engenharia e falhas obrigatórias"""
        if not descricao or not descricao.strip():
            raise ValueError('Descrição da falha é obrigatória')
        problema = {
            'descricao': descricao,
            'severidade': severidade,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'status': 'Aberto'
        }
        self.problemas_engenharia.append(problema)
        self.log_falhas.append(problema)
        return problema

    def registrar_problema_engenharia(self, descricao: str, severidade: str = "Baixa"):
        """Registra problemas de engenharia"""
        return self.registrar_falha(descricao, severidade)


class USG(Equipamento):
    """Ultrassonografia - Vinculado a Emanuel Fernandes Ferreira da Silva Júnior"""
    def __init__(self, fabricante: str = "Fabricante Padrão", pesquisador_responsavel: str = "Emanuel Fernandes Ferreira da Silva Júnior"):
        super().__init__("USG", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.log_analise_imagem = []  # Log de análise de imagem (aeração e diafragma)
        self.treinamentos_confiabilidade = []  # Registros de treinamentos

    def registrar_analise_imagem(self, tipo_analise: str, resultado: str, paciente_id: str = ""):
        """Registra análise de imagem (aeração e diafragma)"""
        analise = {
            'tipo_analise': tipo_analise,  # 'Aeração' ou 'Diafragma'
            'resultado': resultado,
            'paciente_id': paciente_id,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.log_analise_imagem.append(analise)
        self.adicionar_dado_medicao(analise)

    def registrar_treinamento_confiabilidade(self, participante: str, score_confiabilidade: float):
        """Registra treinamento de confiabilidade"""
        treinamento = {
            'participante': participante,
            'score_confiabilidade': score_confiabilidade,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'instrutor': self.pesquisador_responsavel
        }
        self.treinamentos_confiabilidade.append(treinamento)


class EIT(Equipamento):
    """Tomografia de Impedância Elétrica - Vinculado a Maria Eduarda"""
    def __init__(self, fabricante: str = "Fabricante Padrão", pesquisador_responsavel: str = "Maria Eduarda"):
        super().__init__("EIT", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.backup_seguro = []  # Registros de backup seguro

    def padronizar_uso_estudos_clinicos(self, protocolo: str, validacao: str):
        """Padroniza uso em estudos clínicos"""
        padronizacao = {
            'protocolo': protocolo,
            'validacao': validacao,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.adicionar_dado_medicao(padronizacao)

    def registrar_backup_seguro(self, local_backup: str, tamanho_dados: str):
        """Registra backup seguro de dados"""
        backup = {
            'local_backup': local_backup,
            'tamanho_dados': tamanho_dados,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.backup_seguro.append(backup)


class EMG(Equipamento):
    """Eletromiografia de Superfície - Vinculado a Anderson"""
    def __init__(self, fabricante: str = "Fabricante Padrão", pesquisador_responsavel: str = "Anderson"):
        super().__init__("EMG/sEMG", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.protocolos_higienizacao = []  

    def registrar_protocolo_higienizacao(self, equipamento_monitoracao: str, procedimento: str):
        """Registra protocolos de higienização de equipamentos de monitorização avançada"""
        protocolo = {
            'equipamento_monitoracao': equipamento_monitoracao,
            'procedimento': procedimento,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.protocolos_higienizacao.append(protocolo)


class ASL(Equipamento):
    """ASL - Vinculado a Maria Eduarda"""
    def __init__(self, fabricante: str = "Fabricante Padrão", pesquisador_responsavel: str = "Maria Eduarda"):
        super().__init__("ASL", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.tabela_parametros_patologicos = {}  
        self.cenarios_clinicos_registrados = []
        self.datasets_concordancia = []

    def adicionar_parametro_patologico(self, parametro: str, valor_padrao: float, descricao: str):
        """Adiciona parâmetro à tabela de padrões patológicos"""
        self.tabela_parametros_patologicos[parametro] = {
            'valor_padrao': valor_padrao,
            'descricao': descricao,
            'data_cadastro': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }

    def cenarios_clinicos(self, nome_cenario: str = "Insuficiência Respiratória Aguda", descricao: str = "") -> Dict:
        """Configura cenários clínicos para simulação do ASL5000"""
        registro = {
            'nome_cenario': nome_cenario,
            'descricao': descricao or 'Insuficiência Respiratória Aguda',
            'data': datetime.now().strftime('%d/%m/%Y')
        }
        self.cenarios_clinicos_registrados.append(registro)
        return registro

    def concordancia_datasets(self, dataset_simulado: str, banco_clinico: str, concordancia: float) -> Dict:
        """Valida concordância entre dados simulados e bancos clínicos"""
        registro = {
            'dataset_simulado': dataset_simulado,
            'banco_clinico': banco_clinico,
            'concordancia': concordancia,
            'data': datetime.now().strftime('%d/%m/%Y')
        }
        self.datasets_concordancia.append(registro)
        return registro


class PowerLab(Equipamento):
    """PowerLab - Vinculado a Maria Eduarda"""
    def __init__(self, fabricante: str = "ADInstruments", pesquisador_responsavel: str = "Maria Eduarda"):
        super().__init__("PowerLab", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.metas_instrumentacao = []  # Metas de instrumentação
        self.coletas_tcc_ic = []  # Coletas de TCC/Iniciação Científica

    def registrar_meta_instrumentacao(self, meta: str, progresso: float):
        """Registra metas de instrumentação"""
        meta_registro = {
            'meta': meta,
            'progresso': progresso,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.metas_instrumentacao.append(meta_registro)

    def rastrear_instrumentacao(self, status_pops: float = 50.0, status_treinamento: float = 20.0) -> Dict:
        """Monitora progresso de POPs e treinamento prático"""
        registro = {
            'status_pops': status_pops,
            'status_treinamento': status_treinamento,
            'data': datetime.now().strftime('%d/%m/%Y')
        }
        self.metas_instrumentacao.append(registro)
        return registro

    def registrar_coleta_tcc_ic(self, tipo: str, titulo: str, orientador: str):
        """Registra coletas de TCC/Iniciação Científica"""
        coleta = {
            'tipo': tipo,  
            'titulo': titulo,
            'orientador': orientador,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.coletas_tcc_ic.append(coleta)


class MicroFET(Equipamento):
    """MicroFET - Vinculado a Maria Eduarda"""
    def __init__(self, fabricante: str = "Hoggan Scientific", pesquisador_responsavel: str = "Maria Eduarda"):
        super().__init__("microFET", fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.metas_instrumentacao = []  
        self.coletas_tcc_ic = []  

    def registrar_meta_instrumentacao(self, meta: str, progresso: float):
        """Registra metas de instrumentação"""
        meta_registro = {
            'meta': meta,
            'progresso': progresso,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.metas_instrumentacao.append(meta_registro)

    def registrar_coleta_tcc_ic(self, tipo: str, titulo: str, orientador: str):
        """Registra coletas de TCC/Iniciação Científica"""
        coleta = {
            'tipo': tipo,  
            'titulo': titulo,
            'orientador': orientador,
            'data': datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.pesquisador_responsavel
        }
        self.coletas_tcc_ic.append(coleta)


class EquipamentoGenerico(Equipamento):
    """Classe para outros equipamentos não específicos"""
    def __init__(self, nome: str, fabricante: str, pesquisador_responsavel: str = ""):
        super().__init__(nome, fabricante, pesquisador_responsavel=pesquisador_responsavel)
        self.caracteristicas_especificas = {}  


class ModuloGestaoEquipamentos:
    """Módulo de Gestão de Equipamentos e Inovação Tecnológica"""
    def __init__(self):
        self.equipamentos: Dict[str, Equipamento] = {}
        self.perfis_usuarios = {
            'Docente': ['RDA', 'Power Lab C', 'ASL5000', 'Manovacuometro', 'Ventilometro', 'Espirometro'],
            'Aluno': ['Manovacuometro', 'Ventilometro', 'Espirometro']
        }
        self._inicializar_equipamentos()

    def _inicializar_equipamentos(self):
        """Inicializa apenas os equipamentos permitidos pelo modelo atual"""
        rda = RDA()
        self.equipamentos['RDA'] = rda

        powerlab = PowerLab()
        powerlab.nome = 'Power Lab C'
        self.equipamentos['Power Lab C'] = powerlab

        asl = ASL()
        asl.nome = 'ASL5000'
        self.equipamentos['ASL5000'] = asl

        equipamentos_basicos = [
            ('Manovacuometro', 'Fabricante Genérico', 'Maria Eduarda'),
            ('Ventilometro', 'Fabricante Genérico', 'Maria Eduarda'),
            ('Espirometro', 'Fabricante Genérico', 'Maria Eduarda')
        ]

        for nome, fabricante, responsavel in equipamentos_basicos:
            equip = EquipamentoGenerico(nome, fabricante, responsavel)
            if nome in {'Manovacuometro', 'Ventilometro', 'Espirometro'}:
                equip.validar_manutencao('Evento relatado por Maria Eduarda: equipamento em manutenção', status='Em Manutenção')
            self.equipamentos[nome] = equip
            if nome == 'Manovacuometro':
                self.equipamentos['Manovacuômetro'] = equip
            elif nome == 'Ventilometro':
                self.equipamentos['Ventilômetro'] = equip
            elif nome == 'Espirometro':
                self.equipamentos['Espirômetro'] = equip

    def verificar_acesso_equipamento(self, usuario_perfil: str, equipamento_nome: str) -> bool:
        """Verifica se o perfil de usuário tem acesso ao equipamento"""
        if usuario_perfil not in self.perfis_usuarios:
            return False
        return equipamento_nome in self.perfis_usuarios[usuario_perfil]

    def capturar_dados_excel(self, arquivo_excel: str, equipamento_nome: str) -> bool:
        """Captura dados automaticamente de arquivo Excel usando apenas a biblioteca padrão."""
        try:
            if equipamento_nome not in self.equipamentos:
                print(f"✗ Equipamento '{equipamento_nome}' não encontrado")
                return False

            equipamento = self.equipamentos[equipamento_nome]

            if arquivo_excel.lower().endswith('.xlsx'):
                linhas = _ler_planilha_xlsx(arquivo_excel)
            else:
                raise ValueError('Formato de arquivo não suportado. Use .xlsx')

            if not linhas:
                print("✗ Planilha vazia")
                return False

            headers = [str(c) if c is not None else '' for c in linhas[0]]
            for row in linhas[1:]:
                if not row or row[0] is None or row[0] == '':
                    continue
                dados = dict(zip(headers, row))
                equipamento.adicionar_dado_medicao(dados)

            print(f"✓ Dados capturados com sucesso do Excel para {equipamento_nome}")
            return True

        except Exception as e:
            print(f"✗ Erro ao capturar dados do Excel: {e}")
            return False

    def capturar_dados_csv(self, arquivo_csv: str, equipamento_nome: str) -> bool:
        """Captura dados automaticamente de arquivo CSV"""
        try:
            if equipamento_nome not in self.equipamentos:
                print(f"✗ Equipamento '{equipamento_nome}' não encontrado")
                return False

            equipamento = self.equipamentos[equipamento_nome]

            with open(arquivo_csv, 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    equipamento.adicionar_dado_medicao(row)

            print(f"✓ Dados capturados com sucesso do CSV para {equipamento_nome}")
            return True

        except Exception as e:
            print(f"✗ Erro ao capturar dados do CSV: {e}")
            return False

    def importar_dados_equipamento(self, arquivo: str, equipamento_nome: str, comparar_padroes: bool = True) -> bool:
        """Importa dados CSV/XLSX usando pandas e valida contra padrões patológicos do equipamento"""
        try:
            if equipamento_nome not in self.equipamentos:
                print(f"✗ Equipamento '{equipamento_nome}' não encontrado")
                return False
            if pd is None:
                raise ModuleNotFoundError('pandas não disponível')

            equipamento = self.equipamentos[equipamento_nome]
            nome_arquivo = os.path.basename(arquivo).lower()
            if nome_arquivo.endswith('.csv'):
                df = pd.read_csv(arquivo)
            elif nome_arquivo.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(arquivo)
            else:
                raise ValueError('Formato de arquivo não suportado')

            registros = df.to_dict(orient='records')
            equipamento.processar_dados_brutos(registros)

            if comparar_padroes and hasattr(equipamento, 'tabela_parametros_patologicos'):
                for parametro, config in equipamento.tabela_parametros_patologicos.items():
                    if parametro in df.columns and pd.api.types.is_numeric_dtype(df[parametro]):
                        valor_medio = float(df[parametro].mean())
                        desvio = abs(valor_medio - float(config['valor_padrao']))
                        equipamento.adicionar_dado_medicao({
                            'tipo': 'comparacao_padrao',
                            'parametro': parametro,
                            'valor_medio': valor_medio,
                            'valor_padrao': config['valor_padrao'],
                            'desvio': desvio,
                            'descricao': config['descricao']
                        })

            print(f"✓ Dados importados com sucesso para {equipamento_nome} usando pandas")
            return True
        except Exception as e:
            print(f"✗ Erro ao importar dados: {e}")
            return False

    def gerar_pop_mop(self, equipamento_nome: str, tipo: str, arquivo_saida: str = "") -> bool:
        """Gera POP (Procedimento Operacional Padrão) ou MOP automaticamente sem dependências externas."""
        try:
            if equipamento_nome not in self.equipamentos:
                print(f"✗ Equipamento '{equipamento_nome}' não encontrado")
                return False

            equipamento = self.equipamentos[equipamento_nome]

            doc = Document()
            doc.add_heading(f'{tipo} - {equipamento_nome}', 0)

            # Informações básicas
            doc.add_heading('1. Informações Gerais', level=1)
            doc.add_paragraph(f'Equipamento: {equipamento.nome}')
            doc.add_paragraph(f'Fabricante: {equipamento.fabricante}')
            doc.add_paragraph(f'Pesquisador Responsável: {equipamento.pesquisador_responsavel}')
            doc.add_paragraph(f'Última Calibração: {equipamento.ultima_calibracao}')
            doc.add_paragraph(f'Status Operacional: {equipamento.status_operacional}')

            doc.add_heading('2. Procedimentos', level=1)

            if isinstance(equipamento, RDA):
                doc.add_paragraph('2.1 Calibração com Seringa')
                doc.add_paragraph('Procedimento padrão para calibração do Respirador Digital Aberto.')
                doc.add_paragraph('Passos: 1) Preparar seringa, 2) Conectar ao sistema, 3) Realizar medições, 4) Registrar dados.')
                doc.add_paragraph('2.2 Monitoramento de Problemas de Engenharia')
                doc.add_paragraph('Registrar e monitorar problemas identificados durante operação.')

            elif isinstance(equipamento, PowerLab):
                doc.add_paragraph('2.1 Metas de Instrumentação')
                doc.add_paragraph('Objetivos e metas relacionadas à instrumentação.')
                doc.add_paragraph('2.2 Coletas de TCC/Iniciação Científica')
                doc.add_paragraph('Procedimentos para coletas em projetos de TCC e IC.')

            elif isinstance(equipamento, ASL):
                doc.add_paragraph('2.1 Tabela de Parâmetros de Padrões Patológicos')
                doc.add_paragraph('Utilização da tabela desenvolvida pela equipe para análise.')

            else:
                doc.add_paragraph('2.1 Uso operacional do equipamento')
                doc.add_paragraph('Procedimento operacional padrão para uso e manutenção rotineira.')

            # Validação
            doc.add_heading('3. Validação', level=1)
            doc.add_paragraph('Este documento segue os modelos de Validação de Equipamento do LINDEF.')

            if not arquivo_saida:
                arquivo_saida = f"{tipo}_{equipamento_nome.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"

            if not arquivo_saida.lower().endswith('.docx'):
                arquivo_saida = f"{os.path.splitext(arquivo_saida)[0]}.txt"

            doc.save(arquivo_saida)
            print(f"✓ {tipo} gerado com sucesso: {arquivo_saida}")
            return True

        except Exception as e:
            print(f"✗ Erro ao gerar {tipo}: {e}")
            return False

    def validar_meta_equipamento(self, meta_id: str, equipamento_nome: str, pesquisador_nome: str) -> bool:
        """Valida uso de equipamento para meta específica"""
        try:
            if equipamento_nome not in self.equipamentos:
                print(f"✗ Equipamento '{equipamento_nome}' não encontrado")
                return False

            equipamento = self.equipamentos[equipamento_nome]

            
            if "Layane" in pesquisador_nome and equipamento_nome == "RDA" and "Coleta de bancada" in meta_id:
                equipamento.registrar_uso(pesquisador_nome, f"Validação meta {meta_id}")
                print(f"✓ Meta '{meta_id}' validada automaticamente para {pesquisador_nome} usando {equipamento_nome}")
                return True

            
            equipamento.registrar_uso(pesquisador_nome, f"Validação meta {meta_id}")
            print(f"✓ Uso de {equipamento_nome} registrado para validação da meta {meta_id}")
            return True

        except Exception as e:
            print(f"✗ Erro na validação: {e}")
            return False

    def to_dict(self) -> Dict:
        """Converte módulo para dicionário"""
        return {
            'equipamentos': {nome: equip.to_dict() for nome, equip in self.equipamentos.items()},
            'perfis_usuarios': self.perfis_usuarios
        }


class Projeto:
    """Classe representando um projeto de pesquisa vinculado a um pesquisador"""
    def __init__(self, projeto_id: str, titulo: str, area_tematica: str,
                 tecnologias: List[str] = None, orcamento: float = 0.0,
                 data_inicio: str = "", data_termino: str = ""):
        self.projeto_id = projeto_id
        self.titulo = titulo
        self.area_tematica = area_tematica  
        self.tecnologias = tecnologias or []  
        self.orcamento = orcamento
        self.data_inicio = data_inicio
        self.data_termino = data_termino
        self.pesquisadores_associados = []
        self.status = "Ativo"
        self.metas_projeto = []
        self.trl_atual = 1
        self.patentes = 0
        self.pitchs = 0
        self.prototipos = 0
        self.fomento_captado = 0.0

    def adicionar_tecnologia(self, tecnologia: str):
        """Adiciona uma tecnologia ao projeto"""
        if tecnologia not in self.tecnologias:
            self.tecnologias.append(tecnologia)

    def associar_pesquisador(self, pesquisador_nome: str):
        """Associa um pesquisador ao projeto"""
        if pesquisador_nome not in self.pesquisadores_associados:
            self.pesquisadores_associados.append(pesquisador_nome)

    def to_dict(self) -> Dict:
        """Converte projeto para dicionário"""
        return {
            'projeto_id': self.projeto_id,
            'titulo': self.titulo,
            'area_tematica': self.area_tematica,
            'tecnologias': self.tecnologias,
            'orcamento': self.orcamento,
            'data_inicio': self.data_inicio,
            'data_termino': self.data_termino,
            'pesquisadores_associados': self.pesquisadores_associados,
            'status': self.status,
            'trl_atual': self.trl_atual,
            'patentes': self.patentes,
            'pitchs': self.pitchs,
            'prototipos': self.prototipos,
            'fomento_captado': self.fomento_captado
        }


class Meta:
    """Classe representando uma meta de pesquisa"""
    def __init__(self, meta_id: str, descricao: str, meta_prevista: float,
                 alcance_percentual: float = 0, pontos_positivos: str = "",
                 pontos_negativos: int = 0):
        self.meta_id = meta_id
        self.descricao = descricao
        self.meta_prevista = meta_prevista
        self.alcance_percentual = alcance_percentual
        self.pontos_positivos = pontos_positivos
        self.pontos_negativos = pontos_negativos
        self.barreiras = []
        self.solucoes_tomadas = []
        self.link_validacao = ""
        self.dias_lab = 0
        self.status = "Não Validado"  
        self.status_cor = "Vermelho"  
        self.status_exato = ""
        self.data_atualizacao = datetime.now()
        self.barreiras_criticas = []  
        self.justificativa_status = ""

    def validar_status_por_cores(self, alcance: float = None, confirmar_docente: bool = False,
                                  barreiras_pendentes: List[str] = None) -> Tuple[str, str]:
        """
        Valida status baseado em cores da planilha original
        
        Args:
            alcance: Percentual de alcance (opcional)
            confirmar_docente: Se foi confirmado pelo docente
            barreiras_pendentes: Lista de barreiras críticas pendentes
            
        Returns:
            Tupla (status, cor)
        """
        if alcance is not None:
            self.alcance_percentual = alcance
        
        if barreiras_pendentes:
            self.barreiras_criticas = barreiras_pendentes
        
        
        if self.alcance_percentual >= 100 or confirmar_docente:
            self.status = "Validado"
            self.status_cor = "Verde"
            self.justificativa_status = "Alcance 100% ou confirmado pelo docente"
        
        
        elif self.alcance_percentual > 0 or len(self.solucoes_tomadas) > 0:
            self.status = "Validação Parcial"
            self.status_cor = "Amarelo"
            self.justificativa_status = "Meta em andamento ou justificada por barreiras"
        
        
        else:
            self.status = "Não Validado"
            self.status_cor = "Vermelho"
            self.justificativa_status = "0% de alcance ou pendências críticas"
        
        
        if self.barreiras_criticas:
            self.status = "Não Validado"
            self.status_cor = "Vermelho"
            self.justificativa_status = f"Barreiras críticas: {', '.join(self.barreiras_criticas)}"
        
        return (self.status, self.status_cor)

    def atualizar_status(self):
        """Atualiza o status baseado no alcance percentual"""
        self.validar_status_por_cores(self.alcance_percentual)

    def to_dict(self) -> Dict:
        """Converte a meta para dicionário"""
        return {
            'meta_id': self.meta_id,
            'descricao': self.descricao,
            'meta_prevista': self.meta_prevista,
            'alcance_percentual': self.alcance_percentual,
            'pontos_positivos': self.pontos_positivos,
            'pontos_negativos': self.pontos_negativos,
            'barreiras': self.barreiras,
            'barreiras_criticas': self.barreiras_criticas,
            'solucoes_tomadas': self.solucoes_tomadas,
            'link_validacao': self.link_validacao,
            'dias_lab': self.dias_lab,
            'status': self.status,
            'status_cor': self.status_cor,
            'status_exato': self.status_exato,
            'justificativa_status': self.justificativa_status,
            'data_atualizacao': self.data_atualizacao.isoformat()
        }


class Pesquisador:
    """Classe representando um pesquisador no sistema LIMS"""
    def __init__(self, nome: str, titulacao: str, fomento: str = "",
                 programa_pos_grad: str = "", lattes: str = "", orcid: str = "",
                 email_institucional: str = "", metas_previstas: List[Dict[str, any]] = None):
        self.nome = nome
        self.titulacao = titulacao  
        self.fomento = fomento  
        self.programa_pos_grad = programa_pos_grad  
        self.lattes = lattes
        self.orcid = orcid
        self.email_institucional = email_institucional  
        self.metas_previstas = metas_previstas or []  
        self.metas: List[Meta] = []
        self.artigos = []
        self.abstracts = []
        self.revisoes = []
        self.disciplinas_pos_grad = []
        self.atividades_extensao = []
        self.metas_pim_pid = []
        self.projetos: List[Projeto] = []  
        self.tag_financiamento = []  
        self.pops_mops = []  
        self.equipamentos_habilitados = []  
        self.pesquisa_hof = {}  
        self.participacao_eventos = []  
        self.formacao_acadeica = {}  
        self.pontos_positivos_log = []  

    def adicionar_financiamento(self, tag: str):
        """Adiciona tag de financiamento (FACEPE, CAPES, CNPq)"""
        if tag not in self.tag_financiamento:
            self.tag_financiamento.append(tag)

    def vincular_projeto(self, projeto: Projeto):
        """Vincula um projeto ao pesquisador"""
        if projeto not in self.projetos:
            self.projetos.append(projeto)
            projeto.associar_pesquisador(self.nome)

    def registrar_pop_mop(self, tipo: str, equipamento: str, descricao: str, 
                        data_elaboracao: str = "") -> Dict:
        """
        Registra um POP (Procedimento Operacional Padrão) ou MOP
        
        Args:
            tipo: "POP" ou "MOP"
            equipamento: Nome do equipamento (ex: "Ultrassom", "Respirador")
            descricao: Descrição do procedimento
            data_elaboracao: Data de elaboração
            
        Returns:
            Dict com dados do registrofoi atualizando
        """
        doc_id = f"{tipo}_{equipamento.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
        doc = {
            'doc_id': doc_id,
            'tipo': tipo,
            'equipamento': equipamento,
            'descricao': descricao,
            'data_elaboracao': data_elaboracao or datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.nome,
            'status': 'Elaborado'
        }
        self.pops_mops.append(doc)
        return doc

    def habilitar_equipamento(self, equipamento: str, data_habilitacao: str = ""):
        """Registra habilitação em um novo equipamento"""
        habilitacao = {
            'equipamento': equipamento,
            'data_habilitacao': data_habilitacao or datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.nome
        }
        self.equipamentos_habilitados.append(habilitacao)

    def registrar_pesquisa_especifica(self, tipo_pesquisa: str, foco: str, 
                                      progresso: float = 0, observacoes: str = ""):
        """
        Registra pesquisa específica (HOF, TLA, RDA, etc)
        
        Args:
            tipo_pesquisa: Tipo de pesquisa (HOF, TLA, RDA, etc)
            foco: Foco específico
            progresso: % de progresso
            observacoes: Observações gerais
        """
        self.pesquisa_hof[tipo_pesquisa] = {
            'foco': foco,
            'progresso': progresso,
            'observacoes': observacoes,
            'data_registro': datetime.now().isoformat()
        }

    def registrar_participacao_evento(self, evento: str, tipo: str = "NEDTI",
                                      descricao: str = "", data: str = ""):
        """
        Registra participação em eventos/extensão
        
        Args:
            evento: Nome do evento
            tipo: Tipo (NEDTI, Research Day, Oficina, Capacitação)
            descricao: Descrição da participação
            data: Data do evento
        """
        participacao = {
            'evento': evento,
            'tipo': tipo,
            'descricao': descricao,
            'data': data or datetime.now().strftime('%d/%m/%Y'),
            'pesquisador': self.nome
        }
        self.participacao_eventos.append(participacao)
        self.atividades_extensao.append(evento)

    def adicionar_meta(self, meta: Meta):
        """Adiciona uma nova meta ao pesquisador"""
        self.metas.append(meta)

    def adicionar_meta_prevista(self, descricao: str, alcance_percentual: float = 0.0,
                                pontos_positivos: str = "", barreiras: str = "",
                                link_validacao: str = ""):
        """
        Adiciona uma nova meta prevista ao pesquisador
        
        Args:
            descricao: Descrição da meta
            alcance_percentual: Percentual de alcance (0.0 a 100.0)
            pontos_positivos: Pontos positivos ou avanços
            barreiras: Barreiras identificadas
            link_validacao: Link para validação
        """
        meta_dict = {
            "descricao": descricao,
            "alcance_percentual": alcance_percentual,
            "pontos_positivos": pontos_positivos,
            "barreiras": barreiras,
            "link_validacao": link_validacao
        }
        self.metas_previstas.append(meta_dict)

    def registrar_ponto_positivo(self, descricao: str, categoria: str = "Geral", 
                                data: str = None):
        """
        Registra um ponto positivo para análise de produtividade 360º
        
        Args:
            descricao: Descrição do ponto positivo
            categoria: Categoria (Meta, Produção, Evento, etc)
            data: Data do registro (formato DD/MM/YYYY)
        """
        if not data:
            data = datetime.now().strftime('%d/%m/%Y')
        
        ponto = {
            'data': data,
            'categoria': categoria,
            'descricao': descricao
        }
        self.pontos_positivos_log.append(ponto)

    def adicionar_artigo(self, titulo: str, revista: str, data_publicacao: str):
        """Registra um artigo científico"""
        self.artigos.append({
            'titulo': titulo,
            'revista': revista,
            'data_publicacao': data_publicacao
        })

    def adicionar_abstract(self, titulo: str, conferencia: str, data: str):
        """Registra um abstract"""
        self.abstracts.append({
            'titulo': titulo,
            'conferencia': conferencia,
            'data': data
        })

    def adicionar_revisao(self, titulo: str, revista: str, data: str):
        """Registra uma revisão"""
        self.revisoes.append({
            'titulo': titulo,
            'revista': revista,
            'data': data
        })

    def calcular_score_geral(self) -> float:
        """Calcula o score geral baseado em metas e produções"""
        if not self.metas:
            return 0.0
        
        total_alcance = sum(meta.alcance_percentual for meta in self.metas) / len(self.metas)
        producao_weight = len(self.artigos) * 3 + len(self.abstracts) * 2 + len(self.revisoes) * 1.5
        
        return (total_alcance * 0.7) + min((producao_weight / 10) * 0.3 * 100, 100)

    def calcular_score_360(self) -> Dict:
        """Calcula pontos positivos e negativos para produtividade 360º"""
        positivos = []
        negativos = []
        score = 0

        for meta in self.metas:
            if meta.alcance_percentual >= 100:
                score += 10
                positivos.append(f"{meta.descricao}: 100%")
            elif meta.alcance_percentual > 0:
                score += int(meta.alcance_percentual / 10)
                positivos.append(f"{meta.descricao}: {meta.alcance_percentual}%")
            if meta.barreiras:
                score -= len(meta.barreiras)
                negativos.extend([f"{meta.descricao}: {barreira}" for barreira in meta.barreiras])
            if meta.barreiras_criticas:
                score -= len(meta.barreiras_criticas)
                negativos.extend([f"{meta.descricao}: {barreira}" for barreira in meta.barreiras_criticas])

        if self.participacao_eventos:
            score += min(len(self.participacao_eventos), 5)
            positivos.append(f"Participações em eventos: {len(self.participacao_eventos)}")

        return {
            'score_360': score,
            'positivos': positivos,
            'negativos': negativos,
            'score_geral': self.calcular_score_geral()
        }

    def to_dict(self) -> Dict:
        """Converte o pesquisador para dicionário"""
        return {
            'nome': self.nome,
            'titulacao': self.titulacao,
            'fomento': self.fomento,
            'programa_pos_grad': self.programa_pos_grad,
            'lattes': self.lattes,
            'orcid': self.orcid,
            'email_institucional': self.email_institucional,
            'metas_previstas': self.metas_previstas,
            'tag_financiamento': self.tag_financiamento,
            'metas': [meta.to_dict() for meta in self.metas],
            'artigos': self.artigos,
            'abstracts': self.abstracts,
            'revisoes': self.revisoes,
            'projetos': [p.to_dict() for p in self.projetos],
            'pops_mops': self.pops_mops,
            'equipamentos_habilitados': self.equipamentos_habilitados,
            'pesquisa_hof': self.pesquisa_hof,
            'participacao_eventos': self.participacao_eventos,
            'score_geral': self.calcular_score_geral(),
            'score_360': self.calcular_score_360(),
            'pontos_positivos_log': self.pontos_positivos_log
        }


class ModuloREDCap:
    """Integração simulada do REDCap com dicionário e controle de acesso"""
    def __init__(self):
        self.dicionario = {
            'Projeto Layane': {'campo': 'USG Pulmonar', 'status': 'Ativo'},
            'Projeto Maria Eduarda': {'campo': 'RDA', 'status': 'Ativo'}
        }
        self.controle_acesso = {
            'Projeto Layane': ['Layane Santana'],
            'Projeto Maria Eduarda': ['Maria Eduarda Martins']
        }

    def sincronizar_dicionario(self, projeto: str, dados: Dict) -> Dict:
        self.dicionario[projeto] = dados
        return self.dicionario[projeto]

    def verificar_acesso_projeto(self, usuario: str, projeto: str) -> bool:
        return usuario in self.controle_acesso.get(projeto, [])


class SistemaLIMS:
    """Sistema LIMS para monitoramento de pesquisadores e metas"""
    def __init__(self):
        self.pesquisadores: Dict[str, Pesquisador] = {}
        self.projetos: Dict[str, Projeto] = {}
        self.docente_responsavel = None  
        self.modulo_equipamentos = ModuloGestaoEquipamentos()  
        self.modulo_redcap = ModuloREDCap()
        self.calendario_eventos = {
            'CBMI 2026': {'data': '01/10/2026', 'alerta': 'Prazo final do CBMI 2026'},
            'CIF 2026': {'data': '15/11/2026', 'alerta': 'Prazo final do CIF 2026'}
        }

    def atribuir_pontos_360(self, nome_pesquisador: str) -> Dict:
        """Calcula score 360º para um pesquisador"""
        nome_resolvido = self._resolver_nome_pesquisador(nome_pesquisador)
        if not nome_resolvido:
            raise ValueError('Pesquisador não encontrado')
        return self.pesquisadores[nome_resolvido].calcular_score_360()

    def validar_vinculo_meta_equipamento(self, nome_pesquisador: str, meta_id: str) -> Dict:
        """Exige log de uso do RDA e USG quando a meta for 'Coleta Fase 1'"""
        nome_resolvido = self._resolver_nome_pesquisador(nome_pesquisador)
        if not nome_resolvido:
            raise ValueError('Pesquisador não encontrado')

        pesquisador = self.pesquisadores[nome_resolvido]
        meta = next((m for m in pesquisador.metas if m.meta_id == meta_id), None)
        if not meta:
            raise ValueError('Meta não encontrada')

        resultado = {'ok': True, 'requisitos': []}
        if 'Coleta Fase 1' in meta.descricao:
            rda = self.modulo_equipamentos.equipamentos.get('RDA')
            usg = self.modulo_equipamentos.equipamentos.get('USG')
            resultado['requisitos'] = [
                {'equipamento': 'RDA', 'log_registrado': any(uso['proposito'] == f'Validação meta {meta_id}' for uso in getattr(rda, 'historico_uso', []))},
                {'equipamento': 'USG', 'log_registrado': any(uso['proposito'] == f'Validação meta {meta_id}' for uso in getattr(usg, 'historico_uso', []))}
            ]
            resultado['ok'] = all(item['log_registrado'] for item in resultado['requisitos'])
        return resultado

    def importar_dados_equipamento(self, arquivo: str, equipamento_nome: str, comparar_padroes: bool = True) -> bool:
        """Encapsula a importação de dados para o módulo de equipamentos"""
        return self.modulo_equipamentos.importar_dados_equipamento(arquivo, equipamento_nome, comparar_padroes)

    def gerar_dashboard_coleta(self, limite_minimo: int = 100) -> Dict:
        """Gera um resumo de triagem vs elegíveis para dashboard de coleta"""
        anderson = self.pesquisadores.get('Anderson Brasil Xavier')
        maria = self.pesquisadores.get('Maria Eduarda Martins')
        resumo = {
            'pacientes_triados': 0,
            'pacientes_elegiveis': 0,
            'meta_minima': limite_minimo,
            'status': 'Abaixo da meta'
        }
        if anderson:
            resumo['pacientes_triados'] += len(anderson.pesquisa_hof)
        if maria:
            resumo['pacientes_triados'] += len(maria.pesquisa_hof)
        resumo['pacientes_elegiveis'] = max(0, limite_minimo - resumo['pacientes_triados'])
        if resumo['pacientes_triados'] >= limite_minimo:
            resumo['status'] = 'Meta atingida'
        return resumo

    def exportar_documentos_laboratorio(self, pasta_saida: str = "") -> List[str]:
        """Gera documentos de texto para manuais técnicos finalizados sem dependências externas."""
        documentos = [
            ('Avaliação de Diafragma', 'Manual técnico de avaliação de diafragma'),
            ('Aspiração/HOF', 'Manual técnico de aspiração e HOF'),
            ('Protocolo de Higienização', 'Protocolo de higienização do laboratório')
        ]
        arquivos = []
        pasta = pasta_saida or os.getcwd()
        os.makedirs(pasta, exist_ok=True)
        for nome, conteudo in documentos:
            doc = Document()
            doc.add_heading(nome, 0)
            doc.add_paragraph(conteudo)
            arquivo = os.path.join(pasta, f"{nome.replace('/', '_')}.txt")
            doc.save(arquivo)
            arquivos.append(arquivo)
        return arquivos

    def gerar_alertas_eventos(self) -> Dict:
        """Retorna alertas para eventos programados"""
        return self.calendario_eventos

    def _resolver_nome_pesquisador(self, nome_curto: str) -> Optional[str]:
        """Retorna o nome completo do pesquisador com base em nome curto"""
        if not nome_curto:
            return None
        nome_curto = nome_curto.strip()
        if nome_curto in self.pesquisadores:
            return nome_curto

        nome_curto_lower = nome_curto.lower()

        
        correspondencias = [nome for nome in self.pesquisadores
                             if nome.lower().split()[0] == nome_curto_lower]
        if len(correspondencias) == 1:
            return correspondencias[0]

        
        correspondencias = [nome for nome in self.pesquisadores
                             if nome.lower().startswith(nome_curto_lower)]
        if len(correspondencias) == 1:
            return correspondencias[0]

    
        correspondencias = [nome for nome in self.pesquisadores
                             if nome_curto_lower in nome.lower()]
        if len(correspondencias) == 1:
            return correspondencias[0]

        if len(correspondencias) > 1:
            print(f"⚠️ Múltiplas correspondências para '{nome_curto}': {correspondencias}. Use nome completo.")
        return None

    def excluir_pesquisador(self, nome: str) -> bool:
        """Exclui um pesquisador do sistema"""
        nome_resolvido = self._resolver_nome_pesquisador(nome)
        if not nome_resolvido or nome_resolvido not in self.pesquisadores:
            print("✗ Pesquisador não encontrado")
            return False
        
        
        del self.pesquisadores[nome_resolvido]
        print(f"✓ Pesquisador '{nome_resolvido}' excluído com sucesso")
        return True

    def excluir_meta(self, nome_pesquisador: str, meta_id: str) -> bool:
        """Exclui uma meta de um pesquisador"""
        nome_resolvido = self._resolver_nome_pesquisador(nome_pesquisador)
        if not nome_resolvido or nome_resolvido not in self.pesquisadores:
            print("✗ Pesquisador não encontrado")
            return False
        
        pesquisador = self.pesquisadores[nome_resolvido]
        meta_encontrada = None
        for meta in pesquisador.metas:
            if meta.meta_id == meta_id:
                meta_encontrada = meta
                break
        
        if not meta_encontrada:
            print("✗ Meta não encontrada")
            return False
        
        pesquisador.metas.remove(meta_encontrada)
        print(f"✓ Meta '{meta_id}' excluída com sucesso de {nome_resolvido}")
        return True

    def inicializar_lindef(self):
        """
        Initializa o sistema com pesquisadores e estrutura LINDEF reorganizada por categorias acadêmicas
        Integra dados do documento 'Janeiro - LINDEF Monitoramento'
        """
        
        docente_responsavel = Pesquisador(
            nome="Profa. Dra. Shirley Lima Campos",
            titulacao="Professora",
            programa_pos_grad="PPGBAS, PPGFT",
            email_institucional="shirley.campos@ufpe.br",
            lattes="http://lattes.cnpq.br/123456789",  
            orcid="0000-0001-2345-6789"  
        )
        docente_responsavel.adicionar_financiamento("CAPES")
        docente_responsavel.adicionar_financiamento("CNPq")
        self.pesquisadores[docente_responsavel.nome] = docente_responsavel
        self.docente_responsavel = docente_responsavel

        
        lista_docente_responsavel = [docente_responsavel]
        
        
        emanuel = Pesquisador(
            nome="Emanuel Fernandes Ferreira da Silva Júnior",
            titulacao="Doutorando",
            programa_pos_grad="PPGBAS",
            email_institucional="emanuel.fernandes@ufpe.br",
            lattes="http://lattes.cnpq.br/987654321",
            orcid="0000-0002-3456-7890",
            fomento="FACEPE"
        )
        emanuel.adicionar_financiamento("FACEPE")
        metas_emanuel = [
            {"descricao": "Participação na pesquisa HOF (triagem)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Relatório FACEPE", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Fluxograma de triagem TLA", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Fichas de coletas de dados TLA", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Análise de dados pesquisa HOF", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Elaboração de banco de imagens e planilha USG aeração", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Lista de frequência de voluntários e certificados", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Aprimorar conhecimentos (R, pesquisa clínica, instrumentos de avaliação)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Estudo REDCAP", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Ficha de notificação de eventos adversos", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Projeto PIBIC PEEP Table", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_emanuel:
            meta = Meta(
                meta_id=f"EMANUEL_{metas_emanuel.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            emanuel.adicionar_meta(meta)
        emanuel.registrar_pesquisa_especifica("Pesquisa HOF", "Triagem e análise de dados", 60)
        emanuel.registrar_pesquisa_especifica("Pesquisa TLA", "Testes de longa ação", 45)
        emanuel.registrar_pesquisa_especifica("Análise USG", "Análise de imagens de ultrassonografia", 55)
        self.pesquisadores[emanuel.nome] = emanuel

        
        layane = Pesquisador(
            nome="Layane Santana Pereira Costa",
            titulacao="Doutoranda",
            programa_pos_grad="PPGBAS",
            email_institucional="layane.costa@ufpe.br",
            lattes="http://lattes.cnpq.br/112233445",
            orcid="0000-0003-4567-8901",
            fomento="FACEPE"
        )
        layane.adicionar_financiamento("FACEPE")
        metas_layane = [
            {"descricao": "Coleta de bancada", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Correção/Submissão de relatórios PIBICs/PIBITIs FACEPE-CNPq", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Submissão/aprovação de relatórios FACEPE", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Treinamento de equipe", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Coleta saudáveis", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Coleta pacientes", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Resumo ERS", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Submissão ao edital da PROPG (auxílio ERS)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Submissão edital do NEDTI", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Artigo 2", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Análise interina - estudo piloto", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Matrícula em disciplinas", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Artigo RDA - banco de dados", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Treinamento ASL/POWER LAB", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_layane:
            meta = Meta(
                meta_id=f"LAYANE_{metas_layane.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            layane.adicionar_meta(meta)
        layane.registrar_pesquisa_especifica("Coleta de Bancada", "Preparação de amostras e coleta de dados", 75)
        layane.registrar_pesquisa_especifica("Artigo 2", "Desenvolvimento de segundo artigo", 50)
        self.pesquisadores[layane.nome] = layane

        
        jackson = Pesquisador(
            nome="Jackson",
            titulacao="Doutorando",
            programa_pos_grad="PPGBAS",
            email_institucional="jackson.silva@ufpe.br",
            lattes="http://lattes.cnpq.br/556677889",
            orcid="0000-0004-5678-9012",
            fomento="CAPES"
        )
        jackson.adicionar_financiamento("CAPES")
        # Metas para Jackson - Colaboração em tabelas de parâmetros de padrões patológicos ASL
        metas_jackson = [
            {"descricao": "Colaboração em tabelas de parâmetros de padrões patológicos ASL", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_jackson:
            meta = Meta(
                meta_id=f"JACKSON_{metas_jackson.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            jackson.adicionar_meta(meta)
        self.pesquisadores[jackson.nome] = jackson

        
        lista_doutorandos = [emanuel, layane, jackson]
        
        
        
        camilla = Pesquisador(
            nome="Camilla Beatriz Coutinho da Fonsêca",
            titulacao="Mestranda",
            programa_pos_grad="PPGFT",
            email_institucional="camilla.fonseca@ufpe.br",
            lattes="http://lattes.cnpq.br/223344556",
            orcid="0000-0005-6789-0123",
            fomento="FACEPE"
        )
        camilla.adicionar_financiamento("FACEPE")
        metas_camilla = [
            {"descricao": "Qualificação", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Treinamento do protocolo de coleta", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Elaboração das fichas de coleta (equipe RDA)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Revisão do material para início das coletas", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Atualização do artigo SpO2/FiO2 para CHEST", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_camilla:
            meta = Meta(
                meta_id=f"CAMILLA_{metas_camilla.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            camilla.adicionar_meta(meta)
        camilla.registrar_pesquisa_especifica("Projeto RDA", "Pesquisa em Respiração e Dinâmica Aeróbica", 70)
        camilla.participacao_eventos.append({"evento": "NEDTI", "tipo": "Capacitação", "status": "Concluído", "data": datetime.now().strftime('%d/%m/%Y')})
        self.pesquisadores[camilla.nome] = camilla

        
        anderson = Pesquisador(
            nome="Anderson Brasil Xavier",
            titulacao="Mestrando",
            programa_pos_grad="PPGBAS",
            email_institucional="anderson.xavier@ufpe.br",
            lattes="http://lattes.cnpq.br/334455667",
            orcid="0000-0006-7890-1234",
            fomento="CAPES"
        )
        anderson.adicionar_financiamento("CAPES")
        metas_anderson = [
            {"descricao": "Revisão bibliográfica (AET, US, TIE, Programação em R)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Planejamento detalhado do projeto", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Elaboração do Plano de Divulgação e Popularização", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Elaboração do Plano de Gestão de Dados", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Relatórios de prestação de contas", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_anderson:
            meta = Meta(
                meta_id=f"ANDERSON_{metas_anderson.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            anderson.adicionar_meta(meta)
        anderson.registrar_pesquisa_especifica("Revisão Bibliográfica", "AET, US, TIE", 80)
        anderson.registrar_pesquisa_especifica("Estatística com R", "Análise estatística e programação", 65)
        self.pesquisadores[anderson.nome] = anderson
        
        
        maria = Pesquisador(
            nome="Maria Eduarda M. M.",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="maria.eduarda@ufpe.br",
            lattes="http://lattes.cnpq.br/445566778",
            orcid="0000-0007-8901-2345",
            fomento="FACEPE"
        )
        maria.adicionar_financiamento("FACEPE")
        metas_maria = [
            {"descricao": "Treinamento (RDA, Manuvacuometria, Espirometria, etc.)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Confiabilidade do treinamento", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Coleta Fase 1", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Elaboração de POPs", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Instrumentação PowerLab", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Ficha Avaliativa RDA", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Fluxograma de Coleta", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Desenvolvimento Pesquisa RDA", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Recrutamento de Voluntários", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Resumo de Simulação Clínica", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_maria:
            meta = Meta(
                meta_id=f"MARIA_{metas_maria.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            maria.adicionar_meta(meta)
        maria.registrar_pesquisa_especifica("Pesquisa RDA", "Análise de dados respiratórios", 55)
        maria.registrar_pesquisa_especifica("Instrumentação PowerLab", "Calibração e uso de equipamentos", 50)
        maria.participacao_eventos.append({"evento": "Treinamento Clínico", "tipo": "Capacitação", "data": datetime.now().strftime('%d/%m/%Y')})
        self.pesquisadores[maria.nome] = maria

        
        elaine = Pesquisador(
            nome="Elaine A. Souza",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="elaine.souza@ufpe.br",
            lattes="http://lattes.cnpq.br/556677889",
            orcid="0000-0008-9012-3456",
            fomento="CAPES"
        )
        elaine.adicionar_financiamento("CAPES")
        metas_elaine = [
            {"descricao": "Triagem HOF", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Treino de confiabilidade USG", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Análise de dados TIE (Curso R)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Manoscritto C-ARDS", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_elaine:
            meta = Meta(
                meta_id=f"ELAINE_{metas_elaine.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            elaine.adicionar_meta(meta)
        self.pesquisadores[elaine.nome] = elaine

        
        ingrid = Pesquisador(
            nome="Ingrid",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="ingrid.silva@ufpe.br",
            lattes="http://lattes.cnpq.br/667788990",
            orcid="0000-0009-0123-4567",
            fomento="FACEPE"
        )
        ingrid.adicionar_financiamento("FACEPE")
        
        metas_ingrid = [
            {"descricao": "Colaboração no manuscrito de grupo (Anna)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_ingrid:
            meta = Meta(
                meta_id=f"INGRID_{metas_ingrid.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            ingrid.adicionar_meta(meta)
        self.pesquisadores[ingrid.nome] = ingrid

        
        douglas = Pesquisador(
            nome="Douglas",
            titulacao="Mestrando",
            programa_pos_grad="PPGBAS",
            email_institucional="douglas.oliveira@ufpe.br",
            lattes="http://lattes.cnpq.br/778899001",
            orcid="0000-0010-1234-5678",
            fomento="CAPES"
        )
        douglas.adicionar_financiamento("CAPES")
        
        metas_douglas = [
            {"descricao": "Monitoramento acadêmico e atividades de suporte ao laboratório", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_douglas:
            meta = Meta(
                meta_id=f"DOUGLAS_{metas_douglas.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            douglas.adicionar_meta(meta)
        self.pesquisadores[douglas.nome] = douglas

        
        lista_mestrandos = [camilla, anderson, maria, elaine, ingrid, douglas]

        
        
        porfirio = Pesquisador(
            nome="Porfirio",
            titulacao="Pré-Mestrando",
            programa_pos_grad="PPGBAS",
            email_institucional="porfirio.almeida@ufpe.br",
            lattes="http://lattes.cnpq.br/889900112",
            orcid="0000-0011-2345-6789",
            fomento="FACEPE"
        )
        porfirio.adicionar_financiamento("FACEPE")
        
        metas_porfirio = [
            {"descricao": "Foco em Scoping Review de Ultrassonografia (US)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_porfirio:
            meta = Meta(
                meta_id=f"PORFIRIO_{metas_porfirio.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            porfirio.adicionar_meta(meta)
        self.pesquisadores[porfirio.nome] = porfirio

        
        lista_pre_mestrando = [porfirio]

        
    
        lucas = Pesquisador(
            nome="Lucas",
            titulacao="Iniciação Científica",
            programa_pos_grad="PIBIC",
            email_institucional="lucas.santos@ufpe.br",
            lattes="http://lattes.cnpq.br/990011223",
            orcid="0000-0012-3456-7890",
            fomento="PIBIC"
        )
        lucas.adicionar_financiamento("PIBIC")
        
        metas_lucas = [
            {"descricao": "Artigo microFET (TCC)", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Organização de planilhas PIBIC", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Resumos para CONIC/ENEXC", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Submissões ao 13º ISTI", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_lucas:
            meta = Meta(
                meta_id=f"LUCAS_{metas_lucas.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            lucas.adicionar_meta(meta)
        self.pesquisadores[lucas.nome] = lucas

        # Emily
        emily = Pesquisador(
            nome="Emily",
            titulacao="Iniciação Científica",
            programa_pos_grad="PIBIC",
            email_institucional="emily.rodrigues@ufpe.br",
            lattes="http://lattes.cnpq.br/001122334",
            orcid="0000-0013-4567-8901",
            fomento="PIBIC"
        )
        emily.adicionar_financiamento("PIBIC")
        # Metas para Emily - Colaboração em treinamentos de função pulmonar e calibração de equipamentos
        metas_emily = [
            {"descricao": "Colaboração em treinamentos de função pulmonar", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
            {"descricao": "Calibração de equipamentos", "alcance_percentual": 0.0, "pontos_positivos": "", "barreiras": "", "link_validacao": ""},
        ]
        for meta_dict in metas_emily:
            meta = Meta(
                meta_id=f"EMILY_{metas_emily.index(meta_dict)+1}",
                descricao=meta_dict['descricao'],
                meta_prevista=100.0,
                alcance_percentual=meta_dict['alcance_percentual'],
                pontos_positivos=meta_dict['pontos_positivos']
            )
            if meta_dict['barreiras']:
                meta.barreiras.append(meta_dict['barreiras'])
            meta.link_validacao = meta_dict['link_validacao']
            emily.adicionar_meta(meta)
        self.pesquisadores[emily.nome] = emily

        # Lista PIBIC
        lista_pibic = [lucas, emily]

        # Criar Projetos de Referência (Bioengenharia e Inovação Tecnológica)
        projeto_eit = Projeto(
            projeto_id="PROJ_EIT_001",
            titulo="Tomografia de Impedância Elétrica (EIT) - Aplicações em Fisioterapia",
            area_tematica="Bioengenharia e inovação tecnológica",
            tecnologias=["EIT", "Processamento de Sinais", "Biomédica"],
            data_inicio="2023-01-01"
        )
        projeto_eit.associar_pesquisador("Profa. Dra. Shirley Lima Campos")
        self.projetos[projeto_eit.projeto_id] = projeto_eit
        
        projeto_emg = Projeto(
            projeto_id="PROJ_EMG_001",
            titulo="Eletromiografia (EMG) em Avaliação Funcional",
            area_tematica="Bioengenharia e inovação tecnológica",
            tecnologias=["EMG", "Análise de Movimento", "Fisioterapia"],
            data_inicio="2023-01-01"
        )
        projeto_emg.associar_pesquisador("Emanuel Fernandes Ferreira da Silva Júnior")
        self.projetos[projeto_emg.projeto_id] = projeto_emg
        
        projeto_mec = Projeto(
            projeto_id="PROJ_MEC_001",
            titulo="Mecânica Respiratória - Estudos de Função Pulmonar",
            area_tematica="Bioengenharia e inovação tecnológica",
            tecnologias=["PowerLab", "Espirometria", "Análise Respiratória"],
            data_inicio="2023-01-01"
        )
        projeto_mec.associar_pesquisador("Maria Eduarda M. M.")
        projeto_mec.associar_pesquisador("Camilla Beatriz Coutinho da Fonsêca")
        self.projetos[projeto_mec.projeto_id] = projeto_mec
        
        print("✓ Sistema LINDEF inicializado com sucesso - Reorganizado por categorias acadêmicas")
        print(f"  - Docente Responsável: {len(lista_docente_responsavel)}")
        print(f"  - Doutorandos: {len(lista_doutorandos)}")
        print(f"  - Mestrandos: {len(lista_mestrandos)}")
        print(f"  - Pré-Mestrando: {len(lista_pre_mestrando)}")
        print(f"  - PIBIC: {len(lista_pibic)}")
        print(f"  - Total de Pesquisadores: {len(self.pesquisadores)}")
        print(f"  - Projetos de Referência: {len(self.projetos)}")

    def inicializar_lindef_atualizado(self):
        """Inicializa o sistema com os pesquisadores, fomentos e metas atualizados do pedido."""
        self.pesquisadores = {}
        self.projetos = {}
        self.docente_responsavel = None

        docente = Pesquisador(
            nome="Profa. Dra. Shirley Lima Campos",
            titulacao="Professora",
            programa_pos_grad="PPGBAS, PPGFT",
            email_institucional="shirley.campos@ufpe.br",
            lattes="http://lattes.cnpq.br/shirley",
            orcid="0000-0001-0000-0001",
            fomento="PRPG/UFPE"
        )
        docente.adicionar_financiamento("PRPG/UFPE")
        docente.adicionar_financiamento("CNPq")
        docente.adicionar_financiamento("CAPES 001")
        docente.adicionar_financiamento("FACEPE")
        self.pesquisadores[docente.nome] = docente
        self.docente_responsavel = docente

        emanuel = Pesquisador(
            nome="Emanuel Fernandes",
            titulacao="Doutorando",
            programa_pos_grad="PPGBAS",
            email_institucional="emanuel.silvajunior@ufpe.br",
            lattes="http://lattes.cnpq.br/emanuel",
            orcid="0000-0002-0000-0002",
            fomento="FACEPE"
        )
        emanuel.adicionar_financiamento("FACEPE")
        emanuel.adicionar_financiamento("CAPES 001")
        metas_emanuel = [
            ("Triagem HOF", 100.0, "Triagem HOF concluída", [], "100%"),
            ("Fluxograma TLA", 100.0, "Fluxograma TLA concluído", [], "100%"),
            ("Novo Projeto de Caracterização", 100.0, "Projeto de caracterização concluído", [], "100%"),
            ("Artigo SpO2/FiO2 para CHEST", 100.0, "Pronto para submeter", [], "100% - Pronto para submeter"),
            ("Clinical Trials PHENOM-ICU", 95.0, "Em andamento", [], "95%"),
            ("Resumos CBMI 2026", 100.0, "Submetido", [], "100%"),
            ("Análise HOF", 50.0, "Em andamento", [], "50%"),
            ("Estágio Docência", 100.0, "Concluído", [], "100%"),
            ("POP TLA e HOF", 100.0, "Concluído", [], "100%")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_emanuel:
            meta = Meta(meta_id=f"EMANUEL_{len(emanuel.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            if descricao == 'Análise HOF':
                meta.barreiras_criticas.append('Falta de suporte de voluntários')
            emanuel.adicionar_meta(meta)
        emanuel.registrar_pesquisa_especifica('HOF', 'Triagem e análise', 100)
        emanuel.registrar_pesquisa_especifica('TLA', 'Fluxograma', 100)
        emanuel.registrar_pesquisa_especifica('Clinical Trials', 'PHENOM-ICU', 95)
        emanuel.registrar_participacao_evento('CBMI 2026', 'Conferência', 'Resumo concluído', '01/10/2026')
        self.pesquisadores[emanuel.nome] = emanuel

        layane = Pesquisador(
            nome="Layane Santana",
            titulacao="Doutorando",
            programa_pos_grad="PPGBAS",
            email_institucional="layane.santana@ufpe.br",
            lattes="http://lattes.cnpq.br/layane",
            orcid="0000-0003-0000-0003",
            fomento="CAPES 001"
        )
        layane.adicionar_financiamento("CAPES 001")
        layane.adicionar_financiamento("FACEPE")
        metas_layane = [
            ("Coleta de bancada", 50.0, "Em andamento", [], "Corrente"),
            ("Reunião RDA", 50.0, "Em andamento", [], "Corrente"),
            ("Dados USG Pulmonar", 60.0, "Em andamento", [], "60%"),
            ("Artigo 2 Mestrado", 70.0, "Em andamento", [], "70%"),
            ("Scoping Review US", 100.0, "Concluído", [], "100%"),
            ("PIBICS Modelamento", 40.0, "Em andamento", [], "40%"),
            ("Disciplina Aprendizado de Máquina", 100.0, "Concluída", [], "100%")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_layane:
            meta = Meta(meta_id=f"LAYANE_{len(layane.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            layane.adicionar_meta(meta)
        layane.registrar_pesquisa_especifica('USG', 'Dados pulmonares', 60)
        self.pesquisadores[layane.nome] = layane

        jackson = Pesquisador(
            nome="Jackson",
            titulacao="Doutorando",
            programa_pos_grad="PPGBAS",
            email_institucional="jackson@ufpe.br",
            lattes="http://lattes.cnpq.br/jackson",
            orcid="0000-0004-0000-0004",
            fomento="CAPES 001"
        )
        jackson.adicionar_financiamento("CAPES 001")
        self.pesquisadores[jackson.nome] = jackson

        camilla = Pesquisador(
            nome="Camilla Fonseca",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="camilla.fonseca@ufpe.br",
            lattes="http://lattes.cnpq.br/camilla",
            orcid="0000-0005-0000-0005",
            fomento="CAPES 001"
        )
        camilla.adicionar_financiamento("CAPES 001")
        metas_camilla = [
            ("Artigo SpO2/FiO2 CHEST", 100.0, "Concluído", [], "100%"),
            ("Coletas UPAS", 20.0, "Em andamento", [], "20%"),
            ("Artigo Simulação", 20.0, "Em andamento", [], "20%"),
            ("Revisão Sistemática Tecnologias", 15.0, "Em andamento", [], "15%"),
            ("POPs Projeto RDA", 80.0, "Em andamento", [], "80%"),
            ("Qualificação", 100.0, "Concluída", [], "100%"),
            ("Treinamento Protocolo Coleta", 100.0, "Concluído", [], "100%")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_camilla:
            meta = Meta(meta_id=f"CAMILLA_{len(camilla.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            camilla.adicionar_meta(meta)
        self.pesquisadores[camilla.nome] = camilla

        anderson = Pesquisador(
            nome="Anderson Brasil Xavier",
            titulacao="Mestrando",
            programa_pos_grad="PPGBAS",
            email_institucional="anderson.brasil@ufpe.br",
            lattes="http://lattes.cnpq.br/anderson",
            orcid="0000-0006-0000-0006",
            fomento="CAPES 001"
        )
        anderson.adicionar_financiamento("CAPES 001")
        metas_anderson = [
            ("Submissão do Projeto", 100.0, "OK", [], "OK"),
            ("Produtos de Divulgação", 100.0, "OK", [], "OK")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_anderson:
            meta = Meta(meta_id=f"ANDERSON_{len(anderson.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            anderson.adicionar_meta(meta)
        anderson.registrar_pesquisa_especifica('Coleta', 'Meta mínima', 100)
        self.pesquisadores[anderson.nome] = anderson

        maria = Pesquisador(
            nome="Maria Eduarda Martins",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="maria.memm@ufpe.com",
            lattes="http://lattes.cnpq.br/maria",
            orcid="0000-0007-0000-0007",
            fomento="CAPES 001"
        )
        maria.adicionar_financiamento("CAPES 001")
        metas_maria = [
            ("Confiabilidade USG e Treinamento", 50.0, "Em andamento", [], "50%"),
            ("Coleta Fase 1 e UPAS", 10.0, "Em andamento", [], "10%"),
            ("Elaboração POPs Mano e PowerLab", 50.0, "Em andamento", [], "50%"),
            ("Dicionário REDCap", 100.0, "Concluído", [], "100%"),
            ("Ficha Avaliativa RDA", 90.0, "Em andamento", [], "90%")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_maria:
            meta = Meta(meta_id=f"MARIA_{len(maria.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            maria.adicionar_meta(meta)
        maria.registrar_pesquisa_especifica('Coleta Fase 1', 'RDA e USG', 10)
        self.pesquisadores[maria.nome] = maria

        elaine = Pesquisador(
            nome="Elaine de Souza",
            titulacao="Mestranda",
            programa_pos_grad="PPGBAS",
            email_institucional="elaine.asouza@ufpe.br",
            lattes="http://lattes.cnpq.br/elaine",
            orcid="0000-0008-0000-0008",
            fomento="CAPES 001"
        )
        elaine.adicionar_financiamento("CAPES 001")
        self.pesquisadores[elaine.nome] = elaine

        pedro = Pesquisador(
            nome="Pedro Vinicius Porfirio",
            titulacao="Pré-Mestrando",
            programa_pos_grad="PPGBAS",
            email_institucional="pedro.porfirio@ufpe.br",
            lattes="http://lattes.cnpq.br/pedro",
            orcid="0000-0009-0000-0009",
            fomento=""
        )
        metas_pedro = [
            ("Concluir PIBICs", 100.0, "Concluído", [], "100%"),
            ("Poster ATS", 100.0, "Concluído", [], "100%"),
            ("Banco de dados Saudáveis", 60.0, "Em andamento", [], "60%")
        ]
        for descricao, alcance, pontos, barreiras, status_exato in metas_pedro:
            meta = Meta(meta_id=f"PEDRO_{len(pedro.metas)+1}", descricao=descricao, meta_prevista=100.0, alcance_percentual=alcance, pontos_positivos=pontos)
            meta.barreiras.extend(barreiras)
            meta.status_exato = status_exato
            pedro.adicionar_meta(meta)
        self.pesquisadores[pedro.nome] = pedro

        lucas = Pesquisador(
            nome="Lucas",
            titulacao="PIBIC",
            programa_pos_grad="PIBIC",
            email_institucional="lucas@ufpe.br",
            lattes="http://lattes.cnpq.br/lucas",
            orcid="0000-0010-0000-0010",
            fomento="PIBIC"
        )
        lucas.adicionar_financiamento("PIBIC")
        self.pesquisadores[lucas.nome] = lucas

        emily = Pesquisador(
            nome="Emily",
            titulacao="PIBIC",
            programa_pos_grad="PIBIC",
            email_institucional="emily@ufpe.br",
            lattes="http://lattes.cnpq.br/emily",
            orcid="0000-0011-0000-0011",
            fomento="PIBIC"
        )
        emily.adicionar_financiamento("PIBIC")
        self.pesquisadores[emily.nome] = emily

        projeto_rda = Projeto('PROJ_RDA_001', 'Projeto RDA e USG', 'Bioengenharia e pesquisa clínica', ['RDA', 'USG'], 0.0, '2024-01-01')
        projeto_rda.associar_pesquisador(docente.nome)
        projeto_rda.associar_pesquisador(emanuel.nome)
        projeto_rda.associar_pesquisador(layane.nome)
        projeto_rda.associar_pesquisador(maria.nome)
        self.projetos[projeto_rda.projeto_id] = projeto_rda

        projeto_mec = Projeto('PROJ_MEC_001', 'Mecânica Respiratória e Simulação', 'Bioengenharia', ['PowerLab', 'ASL5000', 'Ventilômetro'], 0.0, '2024-01-01')
        projeto_mec.associar_pesquisador(camilla.nome)
        projeto_mec.associar_pesquisador(anderson.nome)
        projeto_mec.associar_pesquisador(maria.nome)
        self.projetos[projeto_mec.projeto_id] = projeto_mec

        print('✓ Sistema LINDEF inicializado com dados atualizados do pedido')
        print(f"  - Docente: {self.docente_responsavel.nome}")
        print(f"  - Pesquisadores: {len(self.pesquisadores)}")
        print(f"  - Projetos: {len(self.projetos)}")

    def importar_dados_legados(self, caminho_arquivo: str) -> bool:
        """
        Importa dados de uma planilha Excel legada com suporte openpyxl
        Lê links de validação e cores de status
        
        Args:
            caminho_arquivo: Caminho do arquivo .xlsx
            
        Returns:
            bool: True se importação bem-sucedida, False caso contrário
        """
        if openpyxl is None:
            print("✗ Funcionalidade indisponível: openpyxl não instalado. Instale com 'pip install openpyxl'.")
            return False

        try:
            workbook = openpyxl.load_workbook(caminho_arquivo)
            worksheet = workbook.active
            
            for idx, row in enumerate(worksheet.iter_rows(min_row=2, values_only=False), start=2):
                try:
                    # Leitura com celulas para obter cores e links
                    nome_cell = row[0]
                    nome = nome_cell.value
                    
                    if not nome:
                        continue
                    
                    titulacao = row[1].value or ""
                    fomento = row[2].value or ""
                    programa = row[3].value or ""
                    lattes = row[4].value or ""
                    orcid = row[5].value or ""
                    email = row[6].value or "" if len(row) > 6 else ""
                    
                    # Criar pesquisador
                    pesquisador = Pesquisador(nome, titulacao, fomento, programa, lattes, orcid, email)
                    
                    # Verificar cor de background (status)
                    if nome_cell.fill and nome_cell.fill.start_color:
                        cor_hex = nome_cell.fill.start_color.rgb
                        print(f"  Cor detectada para {nome}: {cor_hex}")
                    
                    # Tentar extrair links de validação (Google Drive)
                    if len(row) > 7 and row[7].hyperlink:
                        link_validacao = row[7].hyperlink.target
                        pesquisador.metas.append(Meta(
                            meta_id=f"META_{nome.replace(' ', '_')}_001",
                            descricao="Meta importada da planilha",
                            meta_prevista=1.0,
                            alcance_percentual=0
                        ))
                        pesquisador.metas[-1].link_validacao = link_validacao
                    
                    self.pesquisadores[nome] = pesquisador
                    
                except (IndexError, TypeError, AttributeError) as e:
                    print(f"⚠️ Aviso ao processar linha {idx}: {e}")
                    continue
            
            workbook.close()
            print(f"✓ {len(self.pesquisadores)} pesquisadores importados com sucesso")
            print(f"  Planilha processada: {caminho_arquivo}")
            return True
            
        except Exception as e:
            print(f"✗ Erro ao importar dados legados: {e}")
            return False

    def validar_links_google(self, pesquisador_nome: str) -> Dict:
        """
        Valida links Google Drive/Docs presentes nas metas
        
        Args:
            pesquisador_nome: Nome do pesquisador
            
        Returns:
            Dict com status de validação dos links
        """
        nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
        if not nome_resolvido or nome_resolvido not in self.pesquisadores:
            return {'erro': 'Pesquisador não encontrado'}
        
        pesquisador = self.pesquisadores[nome_resolvido]
        links_validacao = []
        problemas = []
        
        for meta in pesquisador.metas:
            if meta.link_validacao:
                # Validar se link é válido
                if "docs.google.com" in meta.link_validacao or "drive.google.com" in meta.link_validacao:
                    links_validacao.append({
                        'meta': meta.meta_id,
                        'link': meta.link_validacao,
                        'status': 'Válido'
                    })
                else:
                    problemas.append(f"Link inválido em {meta.meta_id}")
            else:
                problemas.append(f"Falta Link em {meta.meta_id}")
        
        return {
            'pesquisador': pesquisador_nome,
            'total_metas': len(pesquisador.metas),
            'links_validados': len(links_validacao),
            'links': links_validacao,
            'problemas': problemas
        }

    def registrar_progresso(self, pesquisador_nome: str, meta_id: str,
                           alcance_valor: float, justificativa: str = "",
                           barreiras_criticas: List[str] = None) -> bool:
        """
        Registra o progresso de uma meta com validação de status
        Aceita nome curto (primeiro nome) e resolve para nome completo.
        """
        nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
        if not nome_resolvido:
            print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
            return False

        try:
            if nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{nome_resolvido}' não encontrado")
                return False
            
            pesquisador = self.pesquisadores[nome_resolvido]
            meta = next((m for m in pesquisador.metas if m.meta_id == meta_id), None)
            
            if not meta:
                print(f"✗ Meta '{meta_id}' não encontrada")
                return False
            
            meta.alcance_percentual = min(100, max(0, alcance_valor))
            meta.validar_status_por_cores(alcance=meta.alcance_percentual, 
                                          barreiras_pendentes=barreiras_criticas or [])
            meta.data_atualizacao = datetime.now()
            
            if justificativa:
                meta.dias_lab += 1
                meta.solucoes_tomadas.append({
                    'data': datetime.now().isoformat(),
                    'justificativa': justificativa
                })
                meta.barreiras.append(justificativa)
            
            status_emoji = "🟩" if meta.status_cor == "Verde" else ("🟨" if meta.status_cor == "Amarelo" else "🟥")
            print(f"{status_emoji} Progresso registrado: {meta.descricao} - {meta.alcance_percentual}% [{meta.status}]")
            return True
        except Exception as e:
            print(f"✗ Erro ao registrar progresso: {e}")
            return False
            
        except Exception as e:
            print(f"✗ Erro ao registrar progresso: {e}")
            return False

    def validar_atividades(self, pesquisador_nome: str) -> Dict:
        """
        Valida atividades de um pesquisador com resumo por cores
        
        Args:
            pesquisador_nome: Nome do pesquisador
            
        Returns:
            Dict com resumo de validação por status
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                return {'erro': 'Pesquisador não encontrado'}
            
            pesquisador = self.pesquisadores[nome_resolvido]
            resumo = {
                'Validado': {'quantidade': 0, 'cor': 'Verde'},
                'Validação Parcial': {'quantidade': 0, 'cor': 'Amarelo'},
                'Não Validado': {'quantidade': 0, 'cor': 'Vermelho'},
                'detalhes': []
            }
            
            for meta in pesquisador.metas:
                meta.atualizar_status()
                resumo[meta.status]['quantidade'] += 1
                resumo['detalhes'].append({
                    'meta': meta.descricao,
                    'meta_id': meta.meta_id,
                    'status': meta.status,
                    'cor': meta.status_cor,
                    'alcance': meta.alcance_percentual,
                    'justificativa': meta.justificativa_status
                })
            
            return resumo
            
        except Exception as e:
            print(f"✗ Erro ao validar atividades: {e}")
            return {'erro': str(e)}

    def gerar_indicadores_produtividade(self) -> Dict:
        """
        Gera indicadores gerais de produtividade incluindo projetos e financiamento
        
        Returns:
            Dict com métricas de desempenho
        """
        try:
            total_metas = sum(len(p.metas) for p in self.pesquisadores.values())
            metas_validadas = sum(
                sum(1 for m in p.metas if m.status == "Validado")
                for p in self.pesquisadores.values()
            )
            metas_parciais = sum(
                sum(1 for m in p.metas if m.status == "Validação Parcial")
                for p in self.pesquisadores.values()
            )
            metas_nao_validadas = sum(
                sum(1 for m in p.metas if m.status == "Não Validado")
                for p in self.pesquisadores.values()
            )
            producao_total = sum(
                len(p.artigos) + len(p.abstracts) + len(p.revisoes)
                for p in self.pesquisadores.values()
            )
            
            # Financiamento
            financiamentos_unicos = set()
            for p in self.pesquisadores.values():
                financiamentos_unicos.update(p.tag_financiamento)
            
            # Projetos
            projetos_ativos = sum(1 for p in self.projetos.values() if p.status == "Ativo")
            
            indicadores = {
                'data_geracao': datetime.now().isoformat(),
                'total_pesquisadores': len(self.pesquisadores),
                'total_metas': total_metas,
                'metas_validadas': metas_validadas,
                'metas_parciais': metas_parciais,
                'metas_nao_validadas': metas_nao_validadas,
                'taxa_validacao': (metas_validadas / total_metas * 100) if total_metas > 0 else 0,
                'producao_total': producao_total,
                'projetos_ativos': projetos_ativos,
                'financiamentos': list(financiamentos_unicos),
                'media_score_geral': sum(p.calcular_score_geral() for p in self.pesquisadores.values()) / len(self.pesquisadores) if self.pesquisadores else 0,
                'pesquisadores_top': sorted(
                    [(p.nome, p.calcular_score_geral(), p.tag_financiamento) for p in self.pesquisadores.values()],
                    key=lambda x: x[1],
                    reverse=True
                )
            }
            
            return indicadores
            
        except Exception as e:
            print(f"✗ Erro ao gerar indicadores: {e}")
            return {}

    def gerar_relatorio_indicadores(self) -> None:
        """
        Gera e exibe um relatório formatado dos indicadores de produtividade
        No formato tabular similar ao relatório completo
        """
        try:
            indicadores = self.gerar_indicadores_produtividade()
            
            if not indicadores or 'erro' in indicadores:
                print("✗ Erro ao gerar indicadores de produtividade")
                return
            
            print("\n" + "="*200)
            print("INDICADORES DE PRODUTIVIDADE - SISTEMA LIMS".center(200))
            print("="*200)
            print(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}".center(200))
            print("="*200 + "\n")
            
            # Seção 1: Visão Geral
            print(f"{'─'*200}")
            print("VISÃO GERAL DO SISTEMA".center(200))
            print(f"{'─'*200}")
            print(f"\n┌─ ESTATÍSTICAS GERAIS")
            print(f"│  Total de Pesquisadores: {indicadores['total_pesquisadores']}")
            print(f"│  Total de Metas: {indicadores['total_metas']}")
            print(f"│  Produção Científica Total: {indicadores['producao_total']} (artigos + abstracts + revisões)")
            print(f"│  Projetos Ativos: {indicadores['projetos_ativos']}")
            print(f"│  Financiamentos Disponíveis: {', '.join(indicadores['financiamentos']) if indicadores['financiamentos'] else 'Nenhum'}")
            print(f"│  Score Médio Geral: {indicadores['media_score_geral']:.2f}/100")
            
            # Seção 2: Status das Metas
            print(f"│")
            print(f"├─ STATUS DAS METAS")
            print(f"│  🟩 Validadas: {indicadores['metas_validadas']} ({(indicadores['metas_validadas']/indicadores['total_metas']*100):.1f}%)" if indicadores['total_metas'] > 0 else "│  🟩 Validadas: 0 (0.0%)")
            print(f"│  🟨 Parciais: {indicadores['metas_parciais']} ({(indicadores['metas_parciais']/indicadores['total_metas']*100):.1f}%)" if indicadores['total_metas'] > 0 else "│  🟨 Parciais: 0 (0.0%)")
            print(f"│  🟥 Não Validadas: {indicadores['metas_nao_validadas']} ({(indicadores['metas_nao_validadas']/indicadores['total_metas']*100):.1f}%)" if indicadores['total_metas'] > 0 else "│  🟥 Não Validadas: 0 (0.0%)")
            print(f"│  Taxa Geral de Validação: {indicadores['taxa_validacao']:.1f}%")
            
            # Seção 3: Todos os Pesquisadores
            print(f"│")
            print(f"├─ TODOS OS PESQUISADORES (ordenados por Score Geral)")
            if indicadores['pesquisadores_top']:
                for idx, (nome, score, financiamento) in enumerate(indicadores['pesquisadores_top'], 1):
                    financ_str = f" ({', '.join(financiamento)})" if financiamento else ""
                    print(f"│  {idx:2d}. {nome[:50]:<50} Score: {score:.2f}/100{financ_str}")
            else:
                print(f"│  Nenhum pesquisador encontrado")
            
            print(f"└─")
            print(f"\n{'='*200}\n")
            
        except Exception as e:
            print(f"✗ Erro ao gerar relatório de indicadores: {e}")

    def registrar_pop_mop_laboratorio(self, pesquisador_nome: str, tipo: str,
                                      equipamento: str, descricao: str) -> Dict:
        """
        Registra um POP ou MOP no laboratório
        
        Args:
            pesquisador_nome: Pesquisador responsável
            tipo: "POP" ou "MOP"
            equipamento: Nome do equipamento
            descricao: Descrição do procedimento
            
        Returns:
            Dict com dados do documento registrado
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return {}
            
            pesquisador = self.pesquisadores[nome_resolvido]
            doc = pesquisador.registrar_pop_mop(tipo, equipamento, descricao)
            print(f"✓ {tipo} registrado: {equipamento} - Por {nome_resolvido}")
            return doc
            
        except Exception as e:
            print(f"✗ Erro ao registrar {tipo}: {e}")
            return {}

    def habilitar_pesquisador_equipamento(self, pesquisador_nome: str, 
                                          equipamento: str) -> bool:
        """
        Registra habilitação de um pesquisador em um equipamento
        
        Args:
            pesquisador_nome: Pesquisador
            equipamento: Nome do equipamento (ex: Ultrassom, PowerLab, Respirador)
            
        Returns:
            bool: True se sucesso
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return False
            
            pesquisador = self.pesquisadores[nome_resolvido]
            pesquisador.habilitar_equipamento(equipamento)
            print(f"✓ {nome_resolvido} habilitado em: {equipamento}")
            return True
            
        except Exception as e:
            print(f"✗ Erro ao habilitar equipamento: {e}")
            return False

    def registrar_evento_academico(self, pesquisador_nome: str, evento: str,
                                   tipo: str = "NEDTI", descricao: str = "") -> Dict:
        """
        Registra participação em evento acadêmico (NEDTI, Research Day, etc)
        
        Args:
            pesquisador_nome: Pesquisador
            evento: Nome do evento
            tipo: Tipo de evento (NEDTI, Research Day, Oficina, Capacitação)
            descricao: Descrição da participação
            
        Returns:
            Dict com dados da participação
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return {}
            
            pesquisador = self.pesquisadores[nome_resolvido]
            pesquisador.registrar_participacao_evento(evento, tipo, descricao)
            print(f"✓ Participação registrada: {nome_resolvido} em {evento}")
            return {
                'pesquisador': nome_resolvido,
                'evento': evento,
                'tipo': tipo,
                'data': datetime.now().strftime('%d/%m/%Y')
            }
            
        except Exception as e:
            print(f"✗ Erro ao registrar evento: {e}")
            return {}

    def gerar_relatorio_equipamentos(self) -> Dict:
        """
        Gera relatório de equipamentos e pesquisadores habilitados
        
        Returns:
            Dict com resumo de equipamentos
        """
        try:
            equipamentos_mapeados = {}
            
            for pesquisador in self.pesquisadores.values():
                for equip in pesquisador.equipamentos_habilitados:
                    nome_equip = equip['equipamento']
                    if nome_equip not in equipamentos_mapeados:
                        equipamentos_mapeados[nome_equip] = []
                    equipamentos_mapeados[nome_equip].append({
                        'pesquisador': pesquisador.nome,
                        'data_habilitacao': equip['data_habilitacao']
                    })
            
            return {
                'total_equipamentos_unicos': len(equipamentos_mapeados),
                'equipamentos': equipamentos_mapeados,
                'data_relatorio': datetime.now().isoformat()
            }
            
        except Exception as e:
            print(f"✗ Erro ao gerar relatório: {e}")
            return {}

    def gerar_relatorio_pesquisador(self, pesquisador_nome: str) -> None:
        """
        Gera e exibe um relatório detalhado de um pesquisador específico no terminal
        No formato tabular similar a uma planilha Excel
        
        Args:
            pesquisador_nome: Nome do pesquisador
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return
            
            pesquisador = self.pesquisadores[nome_resolvido]
            
            print("\n" + "="*200)
            print(f"RELATÓRIO INDIVIDUAL - {pesquisador.nome.upper()}".center(200))
            print("="*200)
            print(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}".center(200))
            print("="*200 + "\n")
            
            print(f"{'─'*200}")
            print(f"[1] PESQUISADOR: {pesquisador.nome.upper()}")
            print(f"{'─'*200}")
            
            # Seção 1: Informações Básicas
            print(f"\n┌─ INFORMAÇÕES BÁSICAS")
            print(f"│  Titulação: {pesquisador.titulacao}")
            print(f"│  Email: {pesquisador.email_institucional}")
            print(f"│  Lattes: {pesquisador.lattes if pesquisador.lattes else 'Não informado'}")
            print(f"│  ORCID: {pesquisador.orcid if pesquisador.orcid else 'Não informado'}")
            print(f"│  Programa(s): {pesquisador.programa_pos_grad if pesquisador.programa_pos_grad else 'Não informado'}")
            print(f"│  Fomento: {pesquisador.fomento if pesquisador.fomento else 'Não informado'}")
            print(f"│  Financiamento: {', '.join(pesquisador.tag_financiamento) if pesquisador.tag_financiamento else 'Sem financiamento'}")
            
            # Seção 2: Metas
            print(f"│")
            print(f"├─ METAS ({len(pesquisador.metas)} total)")
            if pesquisador.metas:
                for meta in pesquisador.metas:
                    cor_emoji = "🟩" if meta.status_cor == "Verde" else ("🟨" if meta.status_cor == "Amarelo" else "🟥")
                    print(f"│  {cor_emoji} {meta.descricao}")
                    print(f"│     ID: {meta.meta_id} | Alcance: {meta.alcance_percentual}% | Status: {meta.status}")
                    if meta.link_validacao:
                        print(f"│     Link: {meta.link_validacao}")
                    if meta.pontos_positivos:
                        print(f"│     Pontos Positivos: {meta.pontos_positivos}")
                    if meta.barreiras:
                        print(f"│     Barreiras: {', '.join(meta.barreiras)}")
            else:
                print(f"│  Nenhuma meta registrada")
            
            # Seção 3: Produção Científica
            print(f"│")
            print(f"├─ PRODUÇÃO CIENTÍFICA")
            print(f"│  Artigos: {len(pesquisador.artigos)}")
            if pesquisador.artigos:
                for artigo in pesquisador.artigos[:3]:  # Mostra apenas os 3 primeiros
                    print(f"│    • {artigo['titulo'][:60]}... ({artigo['revista']})")
                if len(pesquisador.artigos) > 3:
                    print(f"│    ... e mais {len(pesquisador.artigos) - 3}")
            
            print(f"│  Abstracts: {len(pesquisador.abstracts)}")
            if pesquisador.abstracts:
                for abstract in pesquisador.abstracts[:2]:
                    print(f"│    • {abstract['titulo'][:50]}... ({abstract['conferencia']})")
                if len(pesquisador.abstracts) > 2:
                    print(f"│    ... e mais {len(pesquisador.abstracts) - 2}")
            
            print(f"│  Revisões: {len(pesquisador.revisoes)}")
            
            # Seção 4: Pesquisas Específicas
            print(f"│")
            print(f"├─ PESQUISAS ESPECÍFICAS")
            if pesquisador.pesquisa_hof:
                for tipo, info in pesquisador.pesquisa_hof.items():
                    print(f"│  • {tipo}: {info['progresso']}% - {info['foco']}")
            else:
                print(f"│  Nenhuma pesquisa específica registrada")
            
            # Seção 5: Equipamentos
            print(f"│")
            print(f"├─ EQUIPAMENTOS HABILITADOS ({len(pesquisador.equipamentos_habilitados)} total)")
            if pesquisador.equipamentos_habilitados:
                for equip in pesquisador.equipamentos_habilitados:
                    print(f"│  ✓ {equip['equipamento']} (Habilitado em: {equip['data_habilitacao']})")
            else:
                print(f"│  Nenhum equipamento registrado")
            
            # Seção 6: POPs/MOPs
            print(f"│")
            print(f"├─ DOCUMENTAÇÃO - POPs/MOPs ({len(pesquisador.pops_mops)} total)")
            if pesquisador.pops_mops:
                for doc in pesquisador.pops_mops:
                    print(f"│  📄 [{doc['tipo']}] {doc['equipamento']} - {doc['data_elaboracao']}")
                    print(f"│     Descrição: {doc['descricao']}")
            else:
                print(f"│  Nenhuma documentação registrada")
            
            # Seção 7: Eventos e Extensão
            print(f"│")
            print(f"├─ EVENTOS E EXTENSÃO ({len(pesquisador.participacao_eventos)} registros)")
            if pesquisador.participacao_eventos:
                for evento in pesquisador.participacao_eventos[:3]:
                    print(f"│  🎓 {evento['evento']} ({evento['tipo']}) - {evento['data']}")
                if len(pesquisador.participacao_eventos) > 3:
                    print(f"│  ... e mais {len(pesquisador.participacao_eventos) - 3}")
            else:
                print(f"│  Nenhum evento registrado")
            
            # Seção 8: Projetos
            print(f"│")
            print(f"├─ PROJETOS ({len(pesquisador.projetos)} associado(s))")
            if pesquisador.projetos:
                for projeto in pesquisador.projetos:
                    print(f"│  • {projeto.titulo}")
                    print(f"│    Tecnologias: {', '.join(projeto.tecnologias)}")
            else:
                print(f"│  Nenhum projeto associado")
            
            # Score Geral
            score = pesquisador.calcular_score_geral()
            print(f"│")
            print(f"└─ SCORE GERAL: {score:.2f}/100")
            
            print(f"\n{'='*200}\n")
            
        except Exception as e:
            print(f"✗ Erro ao gerar relatório do pesquisador: {e}")

    # ==================== MÉTODOS DE INTEGRAÇÃO COM GESTÃO DE EQUIPAMENTOS ====================

    def capturar_dados_equipamento_excel(self, arquivo_excel: str, equipamento_nome: str) -> bool:
        """
        Captura dados automaticamente de arquivo Excel para equipamento específico
        
        Args:
            arquivo_excel: Caminho para o arquivo Excel
            equipamento_nome: Nome do equipamento
            
        Returns:
            bool: Sucesso da operação
        """
        try:
            return self.modulo_equipamentos.capturar_dados_excel(arquivo_excel, equipamento_nome)
        except Exception as e:
            print(f"✗ Erro na captura de dados Excel: {e}")
            return False

    def capturar_dados_equipamento_csv(self, arquivo_csv: str, equipamento_nome: str) -> bool:
        """
        Captura dados automaticamente de arquivo CSV para equipamento específico
        
        Args:
            arquivo_csv: Caminho para o arquivo CSV
            equipamento_nome: Nome do equipamento
            
        Returns:
            bool: Sucesso da operação
        """
        try:
            return self.modulo_equipamentos.capturar_dados_csv(arquivo_csv, equipamento_nome)
        except Exception as e:
            print(f"✗ Erro na captura de dados CSV: {e}")
            return False

    def gerar_documentacao_equipamento(self, equipamento_nome: str, tipo: str, arquivo_saida: str = "") -> bool:
        """
        Gera POP ou MOP para equipamento específico
        
        Args:
            equipamento_nome: Nome do equipamento
            tipo: "POP" ou "MOP"
            arquivo_saida: Caminho opcional para salvar o documento
            
        Returns:
            bool: Sucesso da geração
        """
        try:
            return self.modulo_equipamentos.gerar_pop_mop(equipamento_nome, tipo, arquivo_saida)
        except Exception as e:
            print(f"✗ Erro na geração de documentação: {e}")
            return False

    def validar_acesso_equipamento(self, perfil_usuario: str, equipamento_nome: str) -> bool:
        """
        Verifica se o perfil de usuário tem acesso ao equipamento
        
        Args:
            perfil_usuario: "Docente" ou "Aluno"
            equipamento_nome: Nome do equipamento
            
        Returns:
            bool: Acesso autorizado
        """
        return self.modulo_equipamentos.verificar_acesso_equipamento(perfil_usuario, equipamento_nome)

    def visualizar_registros_equipamentos(self) -> None:
        """
        Visualiza todos os registros de equipamentos organizados por equipamento
        """
        print("\n" + "="*200)
        print("VISUALIZAÇÃO DE REGISTROS DE EQUIPAMENTOS".center(200))
        print("="*200)
        print(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}".center(200))
        print("="*200 + "\n")

        for nome_equip, equipamento in self.modulo_equipamentos.equipamentos.items():
            print(f"{'─'*200}")
            print(f"EQUIPAMENTO: {nome_equip.upper()}")
            print(f"{'─'*200}")

            # Informações básicas
            print(f"\n┌─ INFORMAÇÕES BÁSICAS")
            print(f"│  Fabricante: {equipamento.fabricante}")
            print(f"│  Status Operacional: {equipamento.status_operacional}")
            print(f"│  Última Calibração: {equipamento.ultima_calibracao}")
            print(f"│  Pesquisador Responsável: {equipamento.pesquisador_responsavel}")

            # Histórico de uso
            print(f"│")
            print(f"├─ HISTÓRICO DE USO ({len(equipamento.historico_uso)} registros)")
            if equipamento.historico_uso:
                for uso in equipamento.historico_uso[-5:]:  # Mostra os últimos 5
                    print(f"│  • {uso['data']} - {uso['pesquisador']}: {uso['proposito']}")
                if len(equipamento.historico_uso) > 5:
                    print(f"│  ... e mais {len(equipamento.historico_uso) - 5}")
            else:
                print(f"│  Nenhum registro de uso")

            # Registros específicos por tipo de equipamento
            if isinstance(equipamento, RDA):
                print(f"│")
                print(f"├─ CALIBRAÇÕES COM SERINGA ({len(equipamento.calibracoes_seringa)} registros)")
                if equipamento.calibracoes_seringa:
                    for cal in equipamento.calibracoes_seringa[-3:]:
                        print(f"│  • {cal['data']} - Volume: {cal['volume']}, Erro: {cal['erro_medido']}")
                    if len(equipamento.calibracoes_seringa) > 3:
                        print(f"│  ... e mais {len(equipamento.calibracoes_seringa) - 3}")
                else:
                    print(f"│  Nenhuma calibração registrada")

                print(f"│")
                print(f"├─ PROBLEMAS DE ENGENHARIA ({len(equipamento.problemas_engenharia)} registros)")
                if equipamento.problemas_engenharia:
                    for prob in equipamento.problemas_engenharia[-3:]:
                        print(f"│  • {prob['data']} - {prob['severidade']}: {prob['descricao']}")
                    if len(equipamento.problemas_engenharia) > 3:
                        print(f"│  ... e mais {len(equipamento.problemas_engenharia) - 3}")
                else:
                    print(f"│  Nenhum problema registrado")

            elif isinstance(equipamento, USG):
                print(f"│")
                print(f"├─ ANÁLISES DE IMAGEM ({len(equipamento.analises_imagem)} registros)")
                if equipamento.analises_imagem:
                    for analise in equipamento.analises_imagem[-3:]:
                        print(f"│  • {analise['data']} - {analise['tipo_analise']}: {analise['resultado']}")
                    if len(equipamento.analises_imagem) > 3:
                        print(f"│  ... e mais {len(equipamento.analises_imagem) - 3}")
                else:
                    print(f"│  Nenhuma análise registrada")

                print(f"│")
                print(f"├─ TREINAMENTOS DE CONFIABILIDADE ({len(equipamento.treinamentos_confiabilidade)} registros)")
                if equipamento.treinamentos_confiabilidade:
                    for trein in equipamento.treinamentos_confiabilidade[-3:]:
                        print(f"│  • {trein['data']} - {trein['participante']}: Score {trein['score_confiabilidade']}")
                    if len(equipamento.treinamentos_confiabilidade) > 3:
                        print(f"│  ... e mais {len(equipamento.treinamentos_confiabilidade) - 3}")
                else:
                    print(f"│  Nenhum treinamento registrado")

            elif isinstance(equipamento, EIT):
                print(f"│")
                print(f"├─ PADRONIZAÇÕES PARA ESTUDOS CLÍNICOS ({len(equipamento.padronizacoes_clinicos)} registros)")
                if equipamento.padronizacoes_clinicos:
                    for pad in equipamento.padronizacoes_clinicos[-3:]:
                        print(f"│  • {pad['data']} - Protocolo: {pad['protocolo']}")
                    if len(equipamento.padronizacoes_clinicos) > 3:
                        print(f"│  ... e mais {len(equipamento.padronizacoes_clinicos) - 3}")
                else:
                    print(f"│  Nenhuma padronização registrada")

                print(f"│")
                print(f"├─ BACKUPS SEGUROS ({len(equipamento.backups_seguros)} registros)")
                if equipamento.backups_seguros:
                    for backup in equipamento.backups_seguros[-3:]:
                        print(f"│  • {backup['data']} - {backup['local_backup']}: {backup['tamanho_dados']}")
                    if len(equipamento.backups_seguros) > 3:
                        print(f"│  ... e mais {len(equipamento.backups_seguros) - 3}")
                else:
                    print(f"│  Nenhum backup registrado")

            elif isinstance(equipamento, EMG):
                print(f"│")
                print(f"├─ PROTOCOLOS DE HIGIENIZAÇÃO ({len(equipamento.protocolos_higienizacao)} registros)")
                if equipamento.protocolos_higienizacao:
                    for hig in equipamento.protocolos_higienizacao[-3:]:
                        print(f"│  • {hig['data']} - {hig['equipamento_monitoracao']}: {hig['procedimento']}")
                    if len(equipamento.protocolos_higienizacao) > 3:
                        print(f"│  ... e mais {len(equipamento.protocolos_higienizacao) - 3}")
                else:
                    print(f"│  Nenhum protocolo registrado")

            elif isinstance(equipamento, ASL):
                print(f"│")
                print(f"├─ PARÂMETROS PATOLÓGICOS ({len(equipamento.parametros_patologicos)} registros)")
                if equipamento.parametros_patologicos:
                    for param in equipamento.parametros_patologicos[-3:]:
                        print(f"│  • {param['parametro']}: {param['valor_padrao']} - {param['descricao']}")
                    if len(equipamento.parametros_patologicos) > 3:
                        print(f"│  ... e mais {len(equipamento.parametros_patologicos) - 3}")
                else:
                    print(f"│  Nenhum parâmetro registrado")

            elif isinstance(equipamento, (PowerLab, MicroFET)):
                print(f"│")
                print(f"├─ METAS DE INSTRUMENTAÇÃO ({len(equipamento.metas_instrumentacao)} registros)")
                if equipamento.metas_instrumentacao:
                    for meta in equipamento.metas_instrumentacao[-3:]:
                        print(f"│  • {meta['data']} - {meta['meta']}: {meta['progresso']}%")
                    if len(equipamento.metas_instrumentacao) > 3:
                        print(f"│  ... e mais {len(equipamento.metas_instrumentacao) - 3}")
                else:
                    print(f"│  Nenhuma meta registrada")

                print(f"│")
                print(f"├─ COLETAS TCC/IC ({len(equipamento.coletas_tcc_ic)} registros)")
                if equipamento.coletas_tcc_ic:
                    for coleta in equipamento.coletas_tcc_ic[-3:]:
                        print(f"│  • {coleta['data']} - {coleta['tipo']}: {coleta['titulo']} (Orientador: {coleta['orientador']})")
                    if len(equipamento.coletas_tcc_ic) > 3:
                        print(f"│  ... e mais {len(equipamento.coletas_tcc_ic) - 3}")
                else:
                    print(f"│  Nenhuma coleta registrada")

            print(f"└─")
            print()

        print(f"{'='*200}\n")

    def registrar_equipamento(self, subacao: str, **kwargs) -> bool:
        """
        Registra diferentes operações da aba de gestão de equipamentos em um único ponto.

        Subações:
            uso_meta: Registrar uso de equipamento para meta
            calibracao_rda: Registrar calibração do RDA
            problema_rda: Registrar problema do RDA
            analise_usg: Registrar análise de imagem USG
            treinamento_usg: Registrar treinamento de confiabilidade USG
            backup_eit: Registrar backup seguro da EIT
            higienizacao_emg: Registrar higienização de equipamento EMG
            meta_powerlab: Registrar meta de instrumentação PowerLab
            coleta_powerlab: Registrar coleta TCC/IC PowerLab
        """
        subacao = (subacao or '').strip().lower()
        try:
            if subacao == 'uso_meta':
                meta_id = kwargs.get('meta_id', '')
                equipamento_nome = kwargs.get('equipamento_nome', '')
                pesquisador_nome = kwargs.get('pesquisador_nome', '')
                nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
                if not nome_resolvido:
                    print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                    return False
                return self.modulo_equipamentos.validar_meta_equipamento(meta_id, equipamento_nome, nome_resolvido)

            if subacao == 'calibracao_rda':
                volume = kwargs.get('volume')
                erro_medido = kwargs.get('erro_medido')
                data = kwargs.get('data', '')
                if 'RDA' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento RDA não encontrado")
                    return False
                rda = self.modulo_equipamentos.equipamentos['RDA']
                if not isinstance(rda, RDA):
                    print("✗ Equipamento RDA não é do tipo correto")
                    return False
                rda.registrar_calibracao_seringa(volume, erro_medido, data)
                print(f"✓ Calibração RDA registrada: Volume {volume}, Erro {erro_medido}")
                return True

            if subacao == 'problema_rda':
                descricao = kwargs.get('descricao', '')
                severidade = kwargs.get('severidade', 'Baixa')
                if 'RDA' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento RDA não encontrado")
                    return False
                rda = self.modulo_equipamentos.equipamentos['RDA']
                if not isinstance(rda, RDA):
                    print("✗ Equipamento RDA não é do tipo correto")
                    return False
                rda.registrar_problema_engenharia(descricao, severidade)
                print(f"✓ Problema RDA registrado: {severidade} - {descricao}")
                return True

            if subacao == 'analise_usg':
                tipo_analise = kwargs.get('tipo_analise', '')
                resultado = kwargs.get('resultado', '')
                paciente_id = kwargs.get('paciente_id', '')
                if 'USG' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento USG não encontrado")
                    return False
                usg = self.modulo_equipamentos.equipamentos['USG']
                if not isinstance(usg, USG):
                    print("✗ Equipamento USG não é do tipo correto")
                    return False
                usg.registrar_analise_imagem(tipo_analise, resultado, paciente_id)
                print(f"✓ Análise USG registrada: {tipo_analise} - {resultado}")
                return True

            if subacao == 'treinamento_usg':
                participante = kwargs.get('participante', '')
                score_confiabilidade = kwargs.get('score_confiabilidade')
                if 'USG' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento USG não encontrado")
                    return False
                usg = self.modulo_equipamentos.equipamentos['USG']
                if not isinstance(usg, USG):
                    print("✗ Equipamento USG não é do tipo correto")
                    return False
                usg.registrar_treinamento_confiabilidade(participante, score_confiabilidade)
                print(f"✓ Treinamento USG registrado: {participante} - Score {score_confiabilidade}")
                return True

            if subacao == 'backup_eit':
                local_backup = kwargs.get('local_backup', '')
                tamanho_dados = kwargs.get('tamanho_dados', '')
                if 'EIT' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento EIT não encontrado")
                    return False
                eit = self.modulo_equipamentos.equipamentos['EIT']
                if not isinstance(eit, EIT):
                    print("✗ Equipamento EIT não é do tipo correto")
                    return False
                eit.registrar_backup_seguro(local_backup, tamanho_dados)
                print(f"✓ Backup EIT registrado: {local_backup} - {tamanho_dados}")
                return True

            if subacao == 'higienizacao_emg':
                equipamento_monitoracao = kwargs.get('equipamento_monitoracao', '')
                procedimento = kwargs.get('procedimento', '')
                if 'EMG/sEMG' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento EMG não encontrado")
                    return False
                emg = self.modulo_equipamentos.equipamentos['EMG/sEMG']
                if not isinstance(emg, EMG):
                    print("✗ Equipamento EMG não é do tipo correto")
                    return False
                emg.registrar_protocolo_higienizacao(equipamento_monitoracao, procedimento)
                print(f"✓ Higienização EMG registrada: {equipamento_monitoracao}")
                return True

            if subacao == 'meta_powerlab':
                meta = kwargs.get('meta', '')
                progresso = kwargs.get('progresso')
                if 'PowerLab' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento PowerLab não encontrado")
                    return False
                powerlab = self.modulo_equipamentos.equipamentos['PowerLab']
                if not isinstance(powerlab, PowerLab):
                    print("✗ Equipamento PowerLab não é do tipo correto")
                    return False
                powerlab.registrar_meta_instrumentacao(meta, progresso)
                print(f"✓ Meta PowerLab registrada: {meta} - {progresso}%")
                return True

            if subacao == 'coleta_powerlab':
                tipo = kwargs.get('tipo', '')
                titulo = kwargs.get('titulo', '')
                orientador = kwargs.get('orientador', '')
                if 'PowerLab' not in self.modulo_equipamentos.equipamentos:
                    print("✗ Equipamento PowerLab não encontrado")
                    return False
                powerlab = self.modulo_equipamentos.equipamentos['PowerLab']
                if not isinstance(powerlab, PowerLab):
                    print("✗ Equipamento PowerLab não é do tipo correto")
                    return False
                powerlab.registrar_coleta_tcc_ic(tipo, titulo, orientador)
                print(f"✓ Coleta PowerLab registrada: {tipo} - {titulo}")
                return True

            print(f"✗ Ação inválida para registro de equipamento: {subacao}")
            return False
        except Exception as e:
            print(f"✗ Erro ao registrar operação de equipamento ({subacao}): {e}")
            return False

    def gerar_relatorio_equipamentos_detalhado(self) -> None:
        """
        Gera relatório detalhado de todos os equipamentos e seus dados
        """
        try:
            print("\n" + "="*200)
            print("RELATÓRIO DETALHADO DE EQUIPAMENTOS - SISTEMA LIMS".center(200))
            print("="*200)
            print(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}".center(200))
            print("="*200 + "\n")

            for nome_equip, equipamento in self.modulo_equipamentos.equipamentos.items():
                print(f"{'─'*200}")
                print(f"EQUIPAMENTO: {nome_equip.upper()}")
                print(f"{'─'*200}")

                print(f"\n┌─ INFORMAÇÕES BÁSICAS")
                print(f"│  Fabricante: {equipamento.fabricante}")
                print(f"│  Status Operacional: {equipamento.status_operacional}")
                print(f"│  Última Calibração: {equipamento.ultima_calibracao}")
                print(f"│  Pesquisador Responsável: {equipamento.pesquisador_responsavel}")
                print(f"│  Total de Usos: {len(equipamento.historico_uso)}")
                print(f"│  Total de Dados Medição: {len(equipamento.dados_medicao)}")

                # Informações específicas por tipo de equipamento
                if isinstance(equipamento, RDA):
                    print(f"│  Problemas de Engenharia: {len(equipamento.problemas_engenharia)}")

                elif isinstance(equipamento, USG):
                    print(f"│  Análises de Imagem: {len(equipamento.log_analise_imagem)}")
                    print(f"│  Treinamentos de Confiabilidade: {len(equipamento.treinamentos_confiabilidade)}")

                elif isinstance(equipamento, EIT):
                    print(f"│  Backups Seguros: {len(equipamento.backup_seguro)}")

                elif isinstance(equipamento, EMG):
                    print(f"│  Protocolos de Higienização: {len(equipamento.protocolos_higienizacao)}")

                elif isinstance(equipamento, ASL):
                    print(f"│  Parâmetros Patológicos: {len(equipamento.tabela_parametros_patologicos)}")

                elif isinstance(equipamento, (PowerLab, MicroFET)):
                    print(f"│  Metas de Instrumentação: {len(equipamento.metas_instrumentacao)}")
                    print(f"│  Coletas TCC/IC: {len(equipamento.coletas_tcc_ic)}")

                # Últimos usos
                print(f"│")
                print(f"├─ ÚLTIMOS USOS ({min(5, len(equipamento.historico_uso))} mais recentes)")
                for uso in equipamento.historico_uso[-5:]:
                    print(f"│  • {uso['data']} - {uso['pesquisador']}: {uso['proposito']}")

                # Metas vinculadas
                print(f"│")
                print(f"├─ METAS VINCULADAS ({len(equipamento.metas_vinculadas)} total)")
                for vinculo in equipamento.metas_vinculadas[-3:]:
                    print(f"│  • Meta {vinculo['meta_id']} - {vinculo['pesquisador']} ({vinculo['data_vinculo']})")

                print(f"└─")

            print(f"\n{'='*200}\n")

        except Exception as e:
            print(f"✗ Erro ao gerar relatório detalhado de equipamentos: {e}")

    def to_dict(self) -> Dict:
        """Converte o sistema LIMS completo para dicionário"""
        return {
            'pesquisadores': {nome: p.to_dict() for nome, p in self.pesquisadores.items()},
            'projetos': {nome: p.to_dict() for nome, p in self.projetos.items()},
            'docente_responsavel': self.docente_responsavel.to_dict() if self.docente_responsavel else None,
            'modulo_equipamentos': self.modulo_equipamentos.to_dict()
        }

    def gerar_relatorio_completo(self) -> None:
        try:
            if not self.pesquisadores:
                print("✗ Nenhum pesquisador cadastrado")
                return
            
            print("\n" + "="*200)
            print("RELATÓRIO COMPLETO - TODOS OS PESQUISADORES LINDEF".center(200))
            print("="*200)
            print(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}".center(200))
            print("="*200 + "\n")
            
            # Relatório por pesquisador
            for idx, (nome_pes, pesquisador) in enumerate(self.pesquisadores.items(), 1):
                print(f"\n{'─'*200}")
                print(f"[{idx}] PESQUISADOR: {pesquisador.nome.upper()}")
                print(f"{'─'*200}")
                
                # Seção 1: Informações Básicas
                print(f"\n┌─ INFORMAÇÕES BÁSICAS")
                print(f"│  Titulação: {pesquisador.titulacao}")
                print(f"│  Email: {pesquisador.email_institucional}")
                print(f"│  Lattes: {pesquisador.lattes if pesquisador.lattes else 'Não informado'}")
                print(f"│  ORCID: {pesquisador.orcid if pesquisador.orcid else 'Não informado'}")
                print(f"│  Programa(s): {pesquisador.programa_pos_grad if pesquisador.programa_pos_grad else 'Não informado'}")
                print(f"│  Fomento: {pesquisador.fomento if pesquisador.fomento else 'Não informado'}")
                print(f"│  Financiamento: {', '.join(pesquisador.tag_financiamento) if pesquisador.tag_financiamento else 'Sem financiamento'}")
                
                # Seção 2: Metas
                print(f"│")
                print(f"├─ METAS ({len(pesquisador.metas)} total)")
                if pesquisador.metas:
                    for meta in pesquisador.metas:
                        cor_emoji = "🟩" if meta.status_cor == "Verde" else ("🟨" if meta.status_cor == "Amarelo" else "🟥")
                        print(f"│  {cor_emoji} {meta.descricao}")
                        print(f"│     ID: {meta.meta_id} | Alcance: {meta.alcance_percentual}% | Status: {meta.status}")
                        if meta.link_validacao:
                            print(f"│     Link: {meta.link_validacao}")
                        if meta.pontos_positivos:
                            print(f"│     Pontos Positivos: {meta.pontos_positivos}")
                        if meta.barreiras:
                            print(f"│     Barreiras: {', '.join(meta.barreiras)}")
                else:
                    print(f"│  Nenhuma meta registrada")
                
                # Seção 3: Produção Científica
                print(f"│")
                print(f"├─ PRODUÇÃO CIENTÍFICA")
                print(f"│  Artigos: {len(pesquisador.artigos)}")
                if pesquisador.artigos:
                    for artigo in pesquisador.artigos[:3]:  # Mostra apenas os 3 primeiros
                        print(f"│    • {artigo['titulo'][:60]}... ({artigo['revista']})")
                    if len(pesquisador.artigos) > 3:
                        print(f"│    ... e mais {len(pesquisador.artigos) - 3}")
                
                print(f"│  Abstracts: {len(pesquisador.abstracts)}")
                if pesquisador.abstracts:
                    for abstract in pesquisador.abstracts[:2]:
                        print(f"│    • {abstract['titulo'][:50]}... ({abstract['conferencia']})")
                    if len(pesquisador.abstracts) > 2:
                        print(f"│    ... e mais {len(pesquisador.abstracts) - 2}")
                
                print(f"│  Revisões: {len(pesquisador.revisoes)}")
                
                # Seção 4: Pesquisas Específicas
                print(f"│")
                print(f"├─ PESQUISAS ESPECÍFICAS")
                if pesquisador.pesquisa_hof:
                    for tipo, info in pesquisador.pesquisa_hof.items():
                        print(f"│  • {tipo}: {info['progresso']}% - {info['foco']}")
                else:
                    print(f"│  Nenhuma pesquisa específica registrada")
                
                # Seção 5: Equipamentos
                print(f"│")
                print(f"├─ EQUIPAMENTOS HABILITADOS ({len(pesquisador.equipamentos_habilitados)} total)")
                if pesquisador.equipamentos_habilitados:
                    for equip in pesquisador.equipamentos_habilitados:
                        print(f"│  ✓ {equip['equipamento']} (Habilitado em: {equip['data_habilitacao']})")
                else:
                    print(f"│  Nenhum equipamento registrado")
                
                # Seção 6: POPs/MOPs
                print(f"│")
                print(f"├─ DOCUMENTAÇÃO - POPs/MOPs ({len(pesquisador.pops_mops)} total)")
                if pesquisador.pops_mops:
                    for doc in pesquisador.pops_mops:
                        print(f"│  📄 [{doc['tipo']}] {doc['equipamento']} - {doc['data_elaboracao']}")
                        print(f"│     Descrição: {doc['descricao']}")
                else:
                    print(f"│  Nenhuma documentação registrada")
                
                # Seção 7: Eventos e Extensão
                print(f"│")
                print(f"├─ EVENTOS E EXTENSÃO ({len(pesquisador.participacao_eventos)} registros)")
                if pesquisador.participacao_eventos:
                    for evento in pesquisador.participacao_eventos[:3]:
                        print(f"│  🎓 {evento['evento']} ({evento['tipo']}) - {evento['data']}")
                    if len(pesquisador.participacao_eventos) > 3:
                        print(f"│  ... e mais {len(pesquisador.participacao_eventos) - 3}")
                else:
                    print(f"│  Nenhum evento registrado")
                
                # Seção 8: Projetos
                print(f"│")
                print(f"├─ PROJETOS ({len(pesquisador.projetos)} associado(s))")
                if pesquisador.projetos:
                    for projeto in pesquisador.projetos:
                        print(f"│  • {projeto.titulo}")
                        print(f"│    Tecnologias: {', '.join(projeto.tecnologias)}")
                else:
                    print(f"│  Nenhum projeto associado")
                
                # Score Geral
                score = pesquisador.calcular_score_geral()
                print(f"│")
                print(f"└─ SCORE GERAL: {score:.2f}/100")
            
            # Resumo Final
            print(f"\n{'='*200}")
            print("RESUMO GERAL".center(200))
            print(f"{'='*200}")
            
            total_metas = sum(len(p.metas) for p in self.pesquisadores.values())
            metas_validadas = sum(
                sum(1 for m in p.metas if m.status == "Validado")
                for p in self.pesquisadores.values()
            )
            metas_parciais = sum(
                sum(1 for m in p.metas if m.status == "Validação Parcial")
                for p in self.pesquisadores.values()
            )
            metas_nao_validadas = sum(
                sum(1 for m in p.metas if m.status == "Não Validado")
                for p in self.pesquisadores.values()
            )
            
            total_producao = sum(
                len(p.artigos) + len(p.abstracts) + len(p.revisoes)
                for p in self.pesquisadores.values()
            )
            
            total_equipamentos = len(self.gerar_relatorio_equipamentos()['equipamentos'])
            
            print(f"│ Total de Pesquisadores: {len(self.pesquisadores)}")
            print(f"│ Total de Metas: {total_metas}")
            print(f"│   ├─ 🟩 Validadas: {metas_validadas}")
            print(f"│   ├─ 🟨 Parciais: {metas_parciais}")
            print(f"│   └─ 🟥 Não Validadas: {metas_nao_validadas}")
            print(f"│ Taxa de Validação: {(metas_validadas/total_metas*100):.1f}%" if total_metas > 0 else "│ Taxa de Validação: 0%")
            print(f"│ Produção Científica Total: {total_producao} (artigos + abstracts + revisões)")
            print(f"│ Equipamentos Únicos: {total_equipamentos}")
            print(f"│ Projetos Ativos: {len(self.projetos)}")
            
            media_score = sum(p.calcular_score_geral() for p in self.pesquisadores.values()) / len(self.pesquisadores)
            print(f"│ Score Médio Geral: {media_score:.2f}/100")
            
            print(f"{'='*200}\n")
            
        except Exception as e:
            print(f"✗ Erro ao gerar relatório completo: {e}")

    def exportar_boletim_mensal(self, pesquisador_nome: str, nome_arquivo: str = None) -> bool:
        """
        Exporta boletim mensal em formato Word
        
        Args:
            pesquisador_nome: Nome do pesquisador
            nome_arquivo: Nome do arquivo de saída (opcional)
            
        Returns:
            bool: True se sucesso, False caso contrário
        """
        try:
            nome_resolvido = self._resolver_nome_pesquisador(pesquisador_nome)
            if not nome_resolvido or nome_resolvido not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return False
            
            pesquisador = self.pesquisadores[nome_resolvido]
            
            doc = Document()
            
            # Cabeçalho institucional
            titulo = doc.add_heading('BOLETIM DE DESEMPENHO - LIMS', level=1)
            titulo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            subtitulo = doc.add_paragraph(f"Pesquisador: {pesquisador.nome} ({pesquisador.titulacao})")
            subtitulo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
            doc.add_paragraph(f"Fomento: {pesquisador.fomento}")
            doc.add_paragraph(f"Programa: {pesquisador.programa_pos_grad}")
            doc.add_paragraph(f"LATTES: {pesquisador.lattes}")
            doc.add_paragraph(f"ORCID: {pesquisador.orcid}")
            
            # Metas
            doc.add_heading('METAS DO PERÍODO', level=2)
            if pesquisador.metas:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Meta'
                header_cells[1].text = 'Alcance (%)'
                header_cells[2].text = 'Status'
                header_cells[3].text = 'Pontos (+)'
                header_cells[4].text = 'Pontos (-)'
                
                for meta in pesquisador.metas:
                    row_cells = table.add_row().cells
                    row_cells[0].text = meta.descricao
                    row_cells[1].text = f"{meta.alcance_percentual:.1f}%"
                    row_cells[2].text = meta.status
                    row_cells[3].text = str(meta.pontos_positivos)
                    row_cells[4].text = str(meta.pontos_negativos)
            else:
                doc.add_paragraph("Nenhuma meta registrada.")
            
            # Produção científica
            doc.add_heading('PRODUÇÃO CIENTÍFICA', level=2)
            doc.add_paragraph(f"Artigos: {len(pesquisador.artigos)}")
            doc.add_paragraph(f"Abstracts: {len(pesquisador.abstracts)}")
            doc.add_paragraph(f"Revisões: {len(pesquisador.revisoes)}")
            
            if pesquisador.artigos:
                doc.add_heading('Artigos Publicados:', level=3)
                for artigo in pesquisador.artigos:
                    doc.add_paragraph(f"• {artigo['titulo']} - {artigo['revista']} ({artigo['data_publicacao']})")
            
            # Score geral
            doc.add_heading('RESUMO DE DESEMPENHO', level=2)
            score = pesquisador.calcular_score_geral()
            doc.add_paragraph(f"Score Geral: {score:.2f}/100")
            
            if not nome_arquivo:
                nome_arquivo = f"Boletim_{pesquisador.nome.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
            
            doc.save(nome_arquivo)
            print(f"✓ Boletim exportado: {nome_arquivo}")
            return True
            
        except Exception as e:
            print(f"✗ Erro ao exportar boletim: {e}")
            return False

    def notificar_validacao(self, pesquisador_nome: str, email_pesquisador: str,
                           email_docente: str, smtp_config: Dict = None) -> bool:
        """
        Envia notificação por e-mail
        
        Args:
            pesquisador_nome: Nome do pesquisador
            email_pesquisador: E-mail do pesquisador
            email_docente: E-mail do docente responsável
            smtp_config: Configuração SMTP (host, port, user, password)
            
        Returns:
            bool: True se sucesso, False caso contrário
        """
        try:
            if not smtp_config:
                print("⚠️ Configuração SMTP não fornecida. Notificação não enviada.")
                return False
            
            if pesquisador_nome not in self.pesquisadores:
                print(f"✗ Pesquisador '{pesquisador_nome}' não encontrado")
                return False
            
            pesquisador = self.pesquisadores[pesquisador_nome]
            score = pesquisador.calcular_score_geral()
            
            msg = MIMEMultipart()
            msg['From'] = smtp_config.get('user', '')
            msg['To'] = f"{email_pesquisador}, {email_docente}"
            msg['Subject'] = f"Boletim de Desempenho LIMS - {pesquisador.nome}"
            
            corpo = f"""
Prezado/a,

Segue abaixo o resumo de desempenho do período:

Pesquisador: {pesquisador.nome} ({pesquisador.titulacao})
Score Geral: {score:.2f}/100
Metas: {len(pesquisador.metas)}
Produções: {len(pesquisador.artigos) + len(pesquisador.abstracts) + len(pesquisador.revisoes)}

Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}

Atenciosamente,
Sistema LIMS
            """
            
            msg.attach(MIMEText(corpo, 'plain'))
            
            # Conexão SMTP (comentada por padrão para segurança)
            # try:
            #     server = smtplib.SMTP(smtp_config['host'], smtp_config['port'])
            #     server.starttls()
            #     server.login(smtp_config['user'], smtp_config['password'])
            #     server.send_message(msg)
            #     server.quit()
            #     print(f"✓ Notificação enviada para {email_pesquisador} e {email_docente}")
            #     return True
            # except Exception as e:
            #     print(f"✗ Erro ao enviar e-mail: {e}")
            
            print(f"✓ Notificação preparada (SMTP desabilitado por segurança)")
            return True
            
        except Exception as e:
            print(f"✗ Erro ao notificar: {e}")
            return False

    def padronizar_eit_clinicos(self, protocolo: str, validacao: str) -> bool:
        equipamento = self.modulo_equipamentos.equipamentos.get('EIT')
        if not equipamento:
            equipamento = EIT()
            self.modulo_equipamentos.equipamentos['EIT'] = equipamento
        equipamento.padronizar_uso_estudos_clinicos(protocolo, validacao)
        return True

    def adicionar_parametro_asl(self, parametro: str, valor: float, descricao: str) -> bool:
        equipamento = self.modulo_equipamentos.equipamentos.get('ASL5000')
        if not equipamento:
            return False
        equipamento.adicionar_parametro_patologico(parametro, valor, descricao)
        return True


class PlataformaLINDEF:
    """Camada web persistente para projetos, indicadores e avaliação de uso."""

    def __init__(self, sistema: SistemaLIMS, banco: str = "lindef_platform.db"):
        self.sistema = sistema
        self.banco = banco
        self._criar_banco()
        self._semear_projetos()

    @contextmanager
    def _conectar(self):
        conexao = sqlite3.connect(self.banco)
        conexao.row_factory = sqlite3.Row
        try:
            yield conexao
            conexao.commit()
        except Exception:
            conexao.rollback()
            raise
        finally:
            conexao.close()

    def _criar_banco(self):
        with self._conectar() as conexao:
            conexao.executescript("""
                CREATE TABLE IF NOT EXISTS projetos (
                    id TEXT PRIMARY KEY, titulo TEXT NOT NULL, area TEXT NOT NULL,
                    status TEXT NOT NULL, trl INTEGER NOT NULL DEFAULT 1,
                    patentes INTEGER NOT NULL DEFAULT 0, artigos INTEGER NOT NULL DEFAULT 0,
                    pitchs INTEGER NOT NULL DEFAULT 0, prototipos INTEGER NOT NULL DEFAULT 0,
                    fomento REAL NOT NULL DEFAULT 0, responsavel TEXT NOT NULL DEFAULT ''
                );
                CREATE TABLE IF NOT EXISTS feedbacks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, perfil TEXT NOT NULL,
                    nota INTEGER NOT NULL, comentario TEXT NOT NULL, criado_em TEXT NOT NULL
                );
            """)

    def _semear_projetos(self):
        with self._conectar() as conexao:
            quantidade = conexao.execute("SELECT COUNT(*) FROM projetos").fetchone()[0]
            if quantidade:
                return
            artigos = sum(len(p.artigos) for p in self.sistema.pesquisadores.values())
            for projeto in self.sistema.projetos.values():
                dados = projeto.to_dict()
                conexao.execute(
                    """INSERT OR IGNORE INTO projetos
                    (id, titulo, area, status, trl, artigos, fomento, responsavel)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                    (dados['projeto_id'], dados['titulo'], dados['area_tematica'],
                     dados['status'], dados.get('trl_atual', 1), artigos,
                     dados.get('fomento_captado', dados.get('orcamento', 0)),
                     ', '.join(dados['pesquisadores_associados']))
                )

    def projetos(self):
        with self._conectar() as conexao:
            return [dict(linha) for linha in conexao.execute(
                "SELECT * FROM projetos ORDER BY titulo").fetchall()]

    def dashboard(self):
        indicadores = self.sistema.gerar_indicadores_produtividade()
        projetos = self.projetos()
        return {
            'indicadores': {
                'projetos': len(projetos),
                'projetos_ativos': sum(p['status'] == 'Ativo' for p in projetos),
                'patentes': sum(p['patentes'] for p in projetos),
                'artigos': sum(p['artigos'] for p in projetos),
                'pitchs': sum(p['pitchs'] for p in projetos),
                'prototipos': sum(p['prototipos'] for p in projetos),
                'fomento_captado': sum(p['fomento'] for p in projetos),
                'trl_medio': round(sum(p['trl'] for p in projetos) / len(projetos), 1) if projetos else 0,
                'pesquisadores': indicadores.get('total_pesquisadores', 0),
                'metas_validadas': indicadores.get('metas_validadas', 0)
            },
            'projetos': projetos
        }

    def adicionar_projeto(self, dados):
        obrigatorios = ('id', 'titulo', 'area')
        if any(not str(dados.get(campo, '')).strip() for campo in obrigatorios):
            raise ValueError('id, titulo e area são obrigatórios')
        with self._conectar() as conexao:
            conexao.execute("""INSERT INTO projetos
                (id, titulo, area, status, trl, patentes, artigos, pitchs, prototipos, fomento, responsavel)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""", (
                    dados['id'].strip(), dados['titulo'].strip(), dados['area'].strip(),
                    dados.get('status', 'Ativo'), int(dados.get('trl', 1)),
                    int(dados.get('patentes', 0)), int(dados.get('artigos', 0)),
                    int(dados.get('pitchs', 0)), int(dados.get('prototipos', 0)),
                    float(dados.get('fomento', 0)), dados.get('responsavel', '').strip()))

    def registrar_feedback(self, dados):
        nota = int(dados.get('nota', 0))
        if nota < 1 or nota > 5 or not dados.get('perfil'):
            raise ValueError('perfil e nota de 1 a 5 são obrigatórios')
        with self._conectar() as conexao:
            conexao.execute(
                "INSERT INTO feedbacks (perfil, nota, comentario, criado_em) VALUES (?, ?, ?, ?)",
                (dados['perfil'], nota, dados.get('comentario', '').strip(), datetime.now().isoformat()))

    def executar_comando(self, comando, dados):
        """Executa na web os comandos disponíveis no menu serial."""
        if comando == '2':
            return {'comando': '2', 'subcomandos': ['Adicionar pesquisador', 'Excluir pesquisador'], 'mensagem': 'Selecione uma operação de pesquisadores no menu web.'}
        if comando == '3':
            return {'comando': '3', 'subcomandos': ['Registrar meta', 'Excluir meta'], 'mensagem': 'Selecione uma operação de metas no menu web.'}
        if comando == '14':
            return {'comando': '14', 'subcomandos': ['Uso em meta', 'Calibração RDA', 'Problema RDA', 'Análise USG', 'Treinamento USG', 'Backup EIT', 'Higienização EMG', 'Meta PowerLab', 'Coleta PowerLab'], 'mensagem': 'Selecione a operação específica de equipamento.'}
        if comando in {'0', '24', '31'}:
            if comando == '0':
                self.sistema.inicializar_lindef_atualizado()
            return {'mensagem': 'Comando executado', 'comando': comando}
        if comando in {'13', '17', '19', '26'}:
            return self.sistema.gerar_relatorio_equipamentos()
        if comando in {'18', '25'}:
            return self.sistema.gerar_indicadores_produtividade()
        if comando in {'23', '30'}:
            return {'pesquisadores': list(self.sistema.pesquisadores)}
        if comando in {'22', '29'}:
            return self.sistema.to_dict()
        if comando == '5':
            return self.sistema.validar_atividades(dados.get('pesquisador', ''))
        if comando == '6':
            return self.sistema.validar_links_google(dados.get('pesquisador', ''))
        if comando == '7':
            return self.sistema.registrar_pop_mop_laboratorio(
                dados.get('pesquisador', ''), dados.get('tipo', 'POP'),
                dados.get('equipamento', ''), dados.get('descricao', ''))
        if comando == '8':
            return {'sucesso': self.sistema.habilitar_pesquisador_equipamento(
                dados.get('pesquisador', ''), dados.get('equipamento', ''))}
        if comando == '9':
            return self.sistema.registrar_evento_academico(
                dados.get('pesquisador', ''), dados.get('evento', ''),
                dados.get('tipo', 'NEDTI'), dados.get('descricao', ''))
        if comando == '10':
            return {'sucesso': self.sistema.capturar_dados_equipamento_excel(
                dados.get('arquivo', ''), dados.get('equipamento', ''))}
        if comando == '11':
            return {'sucesso': self.sistema.capturar_dados_equipamento_csv(
                dados.get('arquivo', ''), dados.get('equipamento', ''))}
        if comando == '12':
            return {'sucesso': self.sistema.gerar_documentacao_equipamento(
                dados.get('equipamento', ''), dados.get('tipo', 'POP'), dados.get('arquivo', ''))}
        if comando == '15':
            return {'sucesso': self.sistema.padronizar_eit_clinicos(
                dados.get('protocolo', dados.get('descricao', '')), dados.get('validacao', dados.get('tipo', '')))}
        if comando == '16':
            return {'sucesso': self.sistema.adicionar_parametro_asl(
                dados.get('parametro', dados.get('equipamento', '')), float(dados.get('valor', dados.get('alcance', 0))), dados.get('descricao', ''))}
        if comando == '20' or comando == '27':
            return {'sucesso': self.sistema.exportar_boletim_mensal(
                dados.get('pesquisador', ''), dados.get('arquivo') or None)}
        if comando == '21' or comando == '28':
            nome = self.sistema._resolver_nome_pesquisador(dados.get('pesquisador', ''))
            if not nome:
                return {'erro': 'Pesquisador não encontrado'}
            return self.sistema.pesquisadores[nome].to_dict()
        if comando == '1':
            return {'sucesso': self.sistema.importar_dados_legados(dados.get('arquivo', ''))}
        if comando == '4':
            barreiras = [item.strip() for item in dados.get('barreiras', '').split(',') if item.strip()]
            return {'sucesso': self.sistema.registrar_progresso(
                dados.get('pesquisador', ''), dados.get('meta_id', ''),
                float(dados.get('alcance', 0)), dados.get('justificativa', ''), barreiras)}
        raise ValueError(f'Comando serial {comando} ainda não possui uma ação web.')


PLATAFORMA_HTML = r'''<!doctype html>
<html lang="pt-BR"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>LINDEF | Plataforma de inovação</title>
<style>
:root{--ink:#17221f;--muted:#63716c;--paper:#f5f1e8;--line:#d9d5c9;--green:#1d6b58;--lime:#cce86a;--orange:#ed8b4d}
*{box-sizing:border-box}body{margin:0;color:var(--ink);background:var(--paper);font:15px Georgia,serif}header{padding:34px 6vw 28px;background:var(--ink);color:#fff;display:flex;justify-content:space-between;align-items:end;gap:20px}h1{font-size:clamp(30px,5vw,56px);line-height:.95;margin:0;max-width:570px}header p{color:#d5e2d5;margin:8px 0 0;font:14px Arial,sans-serif}.stamp{font:11px Arial,sans-serif;letter-spacing:1px;color:var(--lime);text-transform:uppercase}.wrap{max-width:1240px;margin:auto;padding:30px 6vw 60px}.grid{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}.card,.panel{background:#fff;border:1px solid var(--line);padding:20px}.card{min-height:130px}.card b{display:block;font:38px Arial,sans-serif;color:var(--green);margin:12px 0 4px}.card span,.label{font:11px Arial,sans-serif;text-transform:uppercase;letter-spacing:1px;color:var(--muted)}.layout{display:grid;grid-template-columns:1.4fr .8fr;gap:18px;margin-top:18px}.panel h2{font-size:22px;margin:0 0 16px}.project{border-top:1px solid var(--line);padding:14px 0;display:grid;grid-template-columns:1fr auto;gap:6px}.project strong{font-size:18px}.project small{display:block;color:var(--muted);font:13px Arial,sans-serif}.badge{background:var(--lime);padding:4px 7px;font:11px Arial,sans-serif;align-self:start}.bar{height:8px;background:#e6e5de;margin:7px 0 14px}.bar i{display:block;height:100%;background:var(--orange)}form{display:grid;gap:9px}input,select,textarea,button{font:14px Arial,sans-serif;padding:10px;border:1px solid var(--line);background:#fff;color:var(--ink)}textarea{min-height:65px}button{background:var(--green);color:white;border:0;cursor:pointer;font-weight:bold}button:hover{background:#134d40}.split{display:grid;grid-template-columns:1fr 1fr;gap:8px}.notice{margin-top:10px;font:12px Arial,sans-serif;color:var(--green)}@media(max-width:800px){header{display:block}.grid{grid-template-columns:1fr 1fr}.layout{grid-template-columns:1fr}}@media(max-width:480px){.grid{grid-template-columns:1fr}.card{min-height:auto}}
</style></head><body><header><div><div class="stamp">LINDEF · inteligência de projetos</div><h1>Pesquisa que avança,<br>evidência que permanece.</h1><p>Painel operacional de inovação, maturidade tecnológica e impacto.</p></div><div class="stamp">Plataforma v1.0</div></header><main class="wrap"><section id="cards" class="grid"></section><section class="layout"><div class="panel"><h2>Projetos e maturidade TRL</h2><div id="projects">Carregando projetos...</div></div><div><div class="panel"><h2>Novo projeto</h2><form id="project-form"><input name="id" placeholder="ID do projeto" required><input name="titulo" placeholder="Título" required><input name="area" placeholder="Área temática" required><div class="split"><input name="trl" type="number" min="1" max="9" value="1" placeholder="TRL"><select name="status"><option>Ativo</option><option>Concluído</option><option>Em pausa</option></select></div><div class="split"><input name="patentes" type="number" min="0" value="0" placeholder="Patentes"><input name="artigos" type="number" min="0" value="0" placeholder="Artigos"></div><div class="split"><input name="pitchs" type="number" min="0" value="0" placeholder="Pitchs"><input name="prototipos" type="number" min="0" value="0" placeholder="Protótipos"></div><input name="fomento" type="number" min="0" step="0.01" value="0" placeholder="Fomento captado (R$)"><input name="responsavel" placeholder="Responsável"><button type="submit">Adicionar ao portfólio</button></form><div id="project-msg" class="notice"></div></div><div class="panel" style="margin-top:18px"><h2>Teste com usuários</h2><form id="feedback-form"><select name="perfil"><option value="bolsista">Bolsista</option><option value="pesquisador">Pesquisador</option><option value="startup">Startup</option><option value="visitante">Visitante</option></select><select name="nota"><option value="5">5 · Excelente</option><option value="4">4 · Bom</option><option value="3">3 · Regular</option><option value="2">2 · Difícil</option><option value="1">1 · Inadequado</option></select><textarea name="comentario" placeholder="Comentário sobre funcionalidade e usabilidade"></textarea><button type="submit">Registrar avaliação</button></form><div id="feedback-msg" class="notice"></div></div></div></section></main><script>
const labels={projetos:'Projetos',projetos_ativos:'Projetos ativos',patentes:'Patentes',artigos:'Artigos',pitchs:'Pitchs',prototipos:'Protótipos',fomento_captado:'Fomento captado',trl_medio:'TRL médio',pesquisadores:'Pesquisadores',metas_validadas:'Metas validadas'};
async function load(){const data=await (await fetch('/api/dashboard')).json();const i=data.indicadores;document.querySelector('#cards').innerHTML=['projetos','patentes','artigos','pitchs','prototipos','fomento_captado','trl_medio','metas_validadas'].map(k=>`<div class="card"><span>${labels[k]}</span><b>${k==='fomento_captado'?'R$ '+i[k].toLocaleString('pt-BR',{maximumFractionDigits:0}):i[k]}</b><span>${k==='trl_medio'?'escala de 1 a 9':'LINDEF'}</span></div>`).join('');document.querySelector('#projects').innerHTML=data.projetos.length?data.projetos.map(p=>`<div class="project"><div><strong>${p.titulo}</strong><small>${p.id} · ${p.area} · ${p.status}</small><small>TRL ${p.trl} · ${p.responsavel||'Equipe LINDEF'}</small><div class="bar"><i style="width:${p.trl/9*100}%"></i></div></div><span class="badge">${p.trl}/9</span></div>`).join(''):'Nenhum projeto cadastrado.'}
async function send(form,url,msg){const body=Object.fromEntries(new FormData(form));const res=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});const data=await res.json();document.querySelector(msg).textContent=data.mensagem||data.erro;if(res.ok){form.reset();load()}}
document.querySelector('#project-form').onsubmit=e=>{e.preventDefault();send(e.target,'/api/projetos','#project-msg')};document.querySelector('#feedback-form').onsubmit=e=>{e.preventDefault();send(e.target,'/api/feedback','#feedback-msg')};load();
</section><section class="panel command-panel"><h2>Central de comandos</h2><p class="command-help">As mesmas operações disponíveis no monitor serial, agora acessíveis pelo navegador.</p><form id="command-form"><select name="comando" id="command-select"><option value="0">0 · Iniciar LIMS</option><option value="1">1 · Importar planilha legada</option><option value="2">2 · Configurar pesquisadores</option><option value="3">3 · Configurar metas</option><option value="4">4 · Atualizar progresso de meta</option><option value="5">5 · Validar atividades</option><option value="6">6 · Validar links Google</option><option value="7">7 · Registrar POP/MOP</option><option value="8">8 · Habilitar pesquisador em equipamento</option><option value="9">9 · Registrar evento</option><option value="10">10 · Capturar equipamento Excel</option><option value="11">11 · Capturar equipamento CSV</option><option value="12">12 · Gerar POP/MOP de equipamento</option><option value="13">13 · Visualizar equipamentos</option><option value="14">14 · Registrar operação de equipamento</option><option value="15">15 · Padronizar EIT clínicos</option><option value="16">16 · Adicionar parâmetro ASL</option><option value="17">17 · Relatório detalhado de equipamentos</option><option value="18">18 · Indicadores de produtividade</option><option value="19">19 · Relatório de equipamentos</option><option value="20">20 · Exportar boletim mensal</option><option value="21">21 · Relatório do pesquisador</option><option value="22">22 · Relatório completo</option><option value="23">23 · Listar pesquisadores</option><option value="24">24 · Encerrar sessão</option><option value="25">25 · Gerar indicadores</option><option value="26">26 · Gerar relatório de equipamentos</option><option value="27">27 · Exportar boletim</option><option value="28">28 · Visualizar pesquisador</option><option value="29">29 · Relatório completo</option><option value="30">30 · Listar nomes completos</option><option value="31">31 · Sair</option></select><input name="pesquisador" placeholder="Pesquisador (quando necessário)"><input name="meta_id" placeholder="ID da meta (quando necessário)"><input name="equipamento" placeholder="Equipamento (quando necessário)"><input name="tipo" placeholder="Tipo: POP, MOP ou evento"><input name="evento" placeholder="Nome do evento"><input name="descricao" placeholder="Descrição"><input name="alcance" type="number" min="0" max="100" placeholder="Alcance (%)"><input name="barreiras" placeholder="Barreiras separadas por vírgula"><input name="arquivo" placeholder="Caminho do arquivo"><button type="submit">Executar comando</button></form><pre id="command-result" class="command-result"></pre></section></main><script>
async function runCommand(form){const body=Object.fromEntries(new FormData(form));const res=await fetch('/api/comandos',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});const data=await res.json();document.querySelector('#command-result').textContent=JSON.stringify(data,null,2)}
document.querySelector('#command-form').onsubmit=e=>{e.preventDefault();runCommand(e.target)};
</script></body></html>'''


def executar_plataforma_web(host='127.0.0.1', porta=8000, banco='lindef_platform.db'):
    sistema = SistemaLIMS()
    sistema.inicializar_lindef_atualizado()
    plataforma = PlataformaLINDEF(sistema, banco)

    class Handler(BaseHTTPRequestHandler):
        def _responder(self, payload, status=200, content_type='application/json; charset=utf-8'):
            conteudo = payload.encode('utf-8') if isinstance(payload, str) else json.dumps(payload, ensure_ascii=False).encode('utf-8')
            self.send_response(status)
            self.send_header('Content-Type', content_type)
            self.send_header('Content-Length', str(len(conteudo)))
            self.end_headers()
            self.wfile.write(conteudo)

        def do_GET(self):
            caminho = urlparse(self.path).path
            if caminho == '/':
                self._responder(PLATAFORMA_HTML, content_type='text/html; charset=utf-8')
            elif caminho == '/api/dashboard':
                self._responder(plataforma.dashboard())
            elif caminho == '/api/projetos':
                self._responder({'projetos': plataforma.projetos()})
            else:
                self._responder({'erro': 'Rota não encontrada'}, 404)

        def do_POST(self):
            caminho = urlparse(self.path).path
            try:
                tamanho = int(self.headers.get('Content-Length', 0))
                dados = json.loads(self.rfile.read(tamanho) or '{}')
                if caminho == '/api/projetos':
                    plataforma.adicionar_projeto(dados)
                    self._responder({'mensagem': 'Projeto adicionado com sucesso'}, 201)
                elif caminho == '/api/feedback':
                    plataforma.registrar_feedback(dados)
                    self._responder({'mensagem': 'Avaliação registrada. Obrigado pela participação.'}, 201)
                elif caminho == '/api/comandos':
                    resultado = plataforma.executar_comando(str(dados.get('comando', '')), dados)
                    self._responder(resultado, 200)
                else:
                    self._responder({'erro': 'Rota não encontrada'}, 404)
            except (ValueError, sqlite3.IntegrityError, json.JSONDecodeError) as erro:
                self._responder({'erro': str(erro)}, 400)

        def log_message(self, *_args):
            return

    servidor = ThreadingHTTPServer((host, porta), Handler)
    print(f'Plataforma LINDEF disponível em http://{host}:{porta}')
    try:
        servidor.serve_forever()
    except KeyboardInterrupt:
        print('\nPlataforma encerrada.')
    finally:
        servidor.server_close()


def menu_interativo():
    """Menu interativo para operações do LIMS"""
    sistema = SistemaLIMS()
    
    while True:
        print("\n" + "="*60)
        print("       LIMS LINDEF")
        print("="*60)
        print("\n[INICIALIZAÇÃO]")
        print("0. Iniciar LIMS")
        print("\n[DADOS BÁSICOS]")
        print("1. Importar dados de planilha legada")
        print("2. Configurar pesquisadores")
        print("3. Configurar metas")
        print("\n[MONITORAMENTO]")
        print("4. Atualizar progresso de meta")
        print("5. Validar atividades")
        print("6. Validar links Google Drive/Docs")
        print("\n[DOCUMENTAÇÃO]")
        print("7. Registrar POP/MOP")
        print("8. Habilitar pesquisador em equipamento")
        print("9. Registrar participação em evento")
        print("\n[GESTÃO DE EQUIPAMENTOS]")
        print("10. Capturar dados de equipamento (Excel)")
        print("11. Capturar dados de equipamento (CSV)")
        print("12. Gerar POP/MOP de equipamento")
        print("13. Visualizar registros de equipamentos")
        print("14. Registrar novo equipamento")
        print("15. Padronizar EIT para estudos clínicos")
        print("16. Adicionar parâmetro ASL")
        print("17. Relatório detalhado de equipamentos")
        print("\n[RELATÓRIOS]")
        print("25. Gerar indicadores de produtividade")
        print("26. Gerar relatório de equipamentos")
        print("27. Exportar boletim mensal")
        print("28. Visualizar dados do pesquisador")
        print("29. Visualizar relatório completo")
        print("\n[SAÍDA]")
        print("30. Listar nomes completos de pesquisadores")
        print("31. Sair")
        print("="*60)
        
        opcao = input("Escolha uma opção: ").strip()
        
        if opcao == "0":
            sistema.inicializar_lindef_atualizado()
        
        elif opcao == "1":
            caminho = input("Caminho do arquivo .xlsx: ").strip()
            sistema.importar_dados_legados(caminho)
        
        elif opcao == "2":
            print("\n--- CONFIGURAR PESQUISADORES ---")
            print("1. Adicionar novo pesquisador")
            print("2. Excluir pesquisador")
            sub_opcao = input("Escolha uma sub-opção: ").strip()
            
            if sub_opcao == "1":
                nome = input("Nome do pesquisador: ").strip()
                titulacao = input("Titulação: ").strip()
                fomento = input("Fomento (FACEPE/CAPES/etc): ").strip()
                programa = input("Programa de Pós-Graduação: ").strip()
                lattes = input("Link Lattes: ").strip()
                orcid = input("ORCID: ").strip()
                email = input("Email (@ufpe.br): ").strip()
                
                pesquisador = Pesquisador(nome, titulacao, fomento, programa, lattes, orcid, email)
                sistema.pesquisadores[nome] = pesquisador
                print(f"✓ Pesquisador '{nome}' adicionado com sucesso")
            
            elif sub_opcao == "2":
                nome = input("Nome do pesquisador a excluir: ").strip()
                sistema.excluir_pesquisador(nome)
            
            else:
                print("✗ Opção inválida")
        
        elif opcao == "3":
            print("\n--- CONFIGURAR METAS ---")
            print("1. Registrar nova meta")
            print("2. Excluir meta")
            sub_opcao = input("Escolha uma sub-opção: ").strip()
            
            if sub_opcao == "1":
                nome = input("Nome do pesquisador: ").strip()
                nome_resolvido = sistema._resolver_nome_pesquisador(nome)
                if not nome_resolvido or nome_resolvido not in sistema.pesquisadores:
                    print("✗ Pesquisador não encontrado")
                    continue
                
                meta_id = input("ID da meta: ").strip()
                descricao = input("Descrição da meta: ").strip()
                try:
                    meta_prevista = float(input("Valor previsto: "))
                    meta = Meta(meta_id, descricao, meta_prevista)
                    sistema.pesquisadores[nome_resolvido].adicionar_meta(meta)
                    print(f"✓ Meta '{descricao}' adicionada para {nome_resolvido}")
                except ValueError:
                    print("✗ Valor inválido")
            
            elif sub_opcao == "2":
                nome = input("Nome do pesquisador: ").strip()
                meta_id = input("ID da meta a excluir: ").strip()
                sistema.excluir_meta(nome, meta_id)
            
            else:
                print("✗ Opção inválida")
        
        elif opcao == "4":
            nome = input("Nome do pesquisador: ").strip()
            meta_id = input("ID da meta: ").strip()
            try:
                alcance = float(input("Alcance percentual (0-100): "))
                justificativa = input("Justificativa (opcional): ").strip()
                barreiras = input("Barreiras críticas separadas por vírgula (opcional): ").strip()
                barreiras_lista = [b.strip() for b in barreiras.split(',')] if barreiras else None
                sistema.registrar_progresso(nome, meta_id, alcance, justificativa, barreiras_lista)
            except ValueError:
                print("✗ Valor inválido")
        
        elif opcao == "5":
            nome = input("Nome do pesquisador: ").strip()
            resultado = sistema.validar_atividades(nome)
            print("\n--- VALIDAÇÃO DE ATIVIDADES (STATUS POR CORES) ---")
            print(json.dumps(resultado, indent=2, ensure_ascii=False))
        
        elif opcao == "6":
            nome = input("Nome do pesquisador: ").strip()
            validacao = sistema.validar_links_google(nome)
            print("\n--- VALIDAÇÃO DE LINKS GOOGLE ---")
            print(json.dumps(validacao, indent=2, ensure_ascii=False))
        
        elif opcao == "7":
            nome = input("Nome do pesquisador: ").strip()
            tipo = input("Tipo (POP/MOP): ").strip().upper()
            if tipo not in ["POP", "MOP"]:
                print("✗ Tipo inválido. Use POP ou MOP")
                continue
            equipamento = input("Equipamento: ").strip()
            descricao = input("Descrição do procedimento: ").strip()
            sistema.registrar_pop_mop_laboratorio(nome, tipo, equipamento, descricao)
        
        elif opcao == "8":
            nome = input("Nome do pesquisador: ").strip()
            equipamento = input("Equipamento (Ultrassom/PowerLab/etc...): ").strip()
            sistema.habilitar_pesquisador_equipamento(nome, equipamento)
        
        elif opcao == "9":
            nome = input("Nome do pesquisador: ").strip()
            evento = input("Nome do evento: ").strip()
            tipo = input("Tipo (NEDTI//Oficina/etc...): ").strip()
            descricao = input("Descrição (opcional): ").strip()
            sistema.registrar_evento_academico(nome, evento, tipo, descricao)
        
        elif opcao == "25":
            sistema.gerar_relatorio_indicadores()
        
        elif opcao == "26":
            relatorio = sistema.gerar_relatorio_equipamentos()
            print("\n--- RELATÓRIO DE EQUIPAMENTOS ---")
            print(json.dumps(relatorio, indent=2, ensure_ascii=False))
        
        elif opcao == "27":
            nome = input("Nome do pesquisador: ").strip()
            arquivo = input("Nome do arquivo (opcional): ").strip()
            sistema.exportar_boletim_mensal(nome, arquivo if arquivo else None)
        
        elif opcao == "28":
            nome = input("Nome do pesquisador: ").strip()
            sistema.gerar_relatorio_pesquisador(nome)
        
        elif opcao == "29":
            sistema.gerar_relatorio_completo()
        
        elif opcao == "30":
            print("\n--- LISTA DE PESQUISADORES CADASTRADOS ---")
            for nome in sistema.pesquisadores:
                print(f"- {nome}")
        
        elif opcao == "31":
            print("Encerrando sistema...")
            break
        
        elif opcao == "10":
            arquivo = input("Caminho do arquivo Excel: ").strip()
            equipamento = input("Nome do equipamento: ").strip()
            sistema.capturar_dados_equipamento_excel(arquivo, equipamento)
        
        elif opcao == "11":
            arquivo = input("Caminho do arquivo CSV: ").strip()
            equipamento = input("Nome do equipamento: ").strip()
            sistema.capturar_dados_equipamento_csv(arquivo, equipamento)
        
        elif opcao == "12":
            equipamento = input("Nome do equipamento: ").strip()
            tipo = input("Tipo (POP/MOP): ").strip().upper()
            if tipo not in ["POP", "MOP"]:
                print("✗ Tipo inválido. Use POP ou MOP")
                continue
            arquivo = input("Nome do arquivo de saída (opcional): ").strip()
            sistema.gerar_documentacao_equipamento(equipamento, tipo, arquivo if arquivo else "")
        
        elif opcao == "13":
            sistema.visualizar_registros_equipamentos()
        
        elif opcao == "14":
            print("\n--- REGISTRAR NOVO REGISTRO DE EQUIPAMENTO ---")
            print("1. Registrar uso de equipamento para meta")
            print("2. Registrar calibração RDA")
            print("3. Registrar problema RDA")
            print("4. Registrar análise USG")
            print("5. Registrar treinamento USG")
            print("6. Registrar backup EIT")
            print("7. Registrar higienização EMG")
            print("8. Registrar meta PowerLab")
            print("9. Registrar coleta PowerLab")
            sub_opcao = input("Escolha uma sub-opção: ").strip()
            
            if sub_opcao == "1":
                meta_id = input("ID da meta: ").strip()
                equipamento = input("Nome do equipamento: ").strip()
                pesquisador = input("Nome do pesquisador: ").strip()
                sistema.registrar_equipamento(
                    'uso_meta',
                    meta_id=meta_id,
                    equipamento_nome=equipamento,
                    pesquisador_nome=pesquisador
                )
            
            elif sub_opcao == "2":
                try:
                    volume = float(input("Volume da calibração: "))
                    erro = float(input("Erro medido: "))
                    data = input("Data (DD/MM/AAAA, opcional): ").strip()
                    sistema.registrar_equipamento(
                        'calibracao_rda',
                        volume=volume,
                        erro_medido=erro,
                        data=data if data else ""
                    )
                except ValueError:
                    print("✗ Valores inválidos")
            
            elif sub_opcao == "3":
                descricao = input("Descrição do problema: ").strip()
                severidade = input("Severidade (Baixa/Média/Alta): ").strip()
                sistema.registrar_equipamento(
                    'problema_rda',
                    descricao=descricao,
                    severidade=severidade
                )
            
            elif sub_opcao == "4":
                tipo = input("Tipo de análise (Aeração/Diafragma): ").strip()
                resultado = input("Resultado: ").strip()
                paciente = input("ID do paciente (opcional): ").strip()
                sistema.registrar_equipamento(
                    'analise_usg',
                    tipo_analise=tipo,
                    resultado=resultado,
                    paciente_id=paciente if paciente else ""
                )
            
            elif sub_opcao == "5":
                participante = input("Nome do participante: ").strip()
                try:
                    score = float(input("Score de confiabilidade (0-100): "))
                    sistema.registrar_equipamento(
                        'treinamento_usg',
                        participante=participante,
                        score_confiabilidade=score
                    )
                except ValueError:
                    print("✗ Score inválido")
            
            elif sub_opcao == "6":
                local = input("Local do backup: ").strip()
                tamanho = input("Tamanho dos dados: ").strip()
                sistema.registrar_equipamento(
                    'backup_eit',
                    local_backup=local,
                    tamanho_dados=tamanho
                )
            
            elif sub_opcao == "7":
                equipamento_mon = input("Equipamento de monitoração: ").strip()
                procedimento = input("Procedimento de higienização: ").strip()
                sistema.registrar_equipamento(
                    'higienizacao_emg',
                    equipamento_monitoracao=equipamento_mon,
                    procedimento=procedimento
                )
            
            elif sub_opcao == "8":
                meta = input("Descrição da meta: ").strip()
                try:
                    progresso = float(input("Progresso (%): "))
                    sistema.registrar_equipamento(
                        'meta_powerlab',
                        meta=meta,
                        progresso=progresso
                    )
                except ValueError:
                    print("✗ Progresso inválido")
            
            elif sub_opcao == "9":
                tipo = input("Tipo (TCC/Iniciação Científica/etc...): ").strip()
                titulo = input("Título: ").strip()
                orientador = input("Orientador: ").strip()
                sistema.registrar_equipamento(
                    'coleta_powerlab',
                    tipo=tipo,
                    titulo=titulo,
                    orientador=orientador
                )
            
            else:
                print("✗ Opção inválida")
        
        elif opcao == "15":
            protocolo = input("Protocolo de padronização: ").strip()
            validacao = input("Validação: ").strip()
            sistema.padronizar_eit_clinicos(protocolo, validacao)
        
        elif opcao == "16":
            parametro = input("Nome do parâmetro: ").strip()
            try:
                valor = float(input("Valor padrão: "))
                descricao = input("Descrição: ").strip()
                sistema.adicionar_parametro_asl(parametro, valor, descricao)
            except ValueError:
                print("✗ Valor inválido")
        
        elif opcao == "17":
            sistema.gerar_relatorio_equipamentos_detalhado()
        
        elif opcao == "18":
            sistema.gerar_indicadores_produtividade()
        
        elif opcao == "19":
            relatorio = sistema.gerar_relatorio_equipamentos()
            print("\n--- RELATÓRIO DE EQUIPAMENTOS ---")
            print(json.dumps(relatorio, indent=2, ensure_ascii=False))
        
        elif opcao == "20":
            nome = input("Nome do pesquisador: ").strip()
            arquivo = input("Nome do arquivo (opcional): ").strip()
            sistema.exportar_boletim_mensal(nome, arquivo if arquivo else None)
        
        elif opcao == "21":
            nome = input("Nome do pesquisador: ").strip()
            sistema.gerar_relatorio_pesquisador(nome)
        
        elif opcao == "22":
            sistema.gerar_relatorio_completo()
        
        elif opcao == "23":
            print("\n--- LISTA DE PESQUISADORES CADASTRADOS ---")
            for nome in sistema.pesquisadores:
                print(f"- {nome}")
        
        elif opcao == "24":
            print("Encerrando sistema...")
            break
        
        else:
            print("✗ Opção inválida")


if __name__ == "__main__":
    if '--web' in sys.argv:
        host = os.getenv('LIS_HOST', '0.0.0.0')
        porta = int(os.getenv('PORT', '8000'))
        banco = os.getenv('LIS_DATABASE', 'lindef_platform.db')
        executar_plataforma_web(host, porta, banco)
    else:
        menu_interativo()
